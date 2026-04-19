"""Generate a single-player USTA scouting report from tennisrecord.com data.

Usage:
    python3 skills/player-scout/player_report.py --first FIRSTNAME --last LASTNAME \\
        --location "CITY, STATE"

Writes a Word (.docx) file to the workspace-local ``reports/`` folder.
"""

from __future__ import annotations

import argparse
import re
import subprocess
import sys
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from urllib.parse import quote

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_URL = "https://www.tennisrecord.com"
SEARCH_URL = f"{BASE_URL}/adult/search.aspx"

# ---------------------------------------------------------------------------
# Nickname / alias lookup table  (2.1)
# ---------------------------------------------------------------------------

NICKNAME_MAP: Dict[str, List[str]] = {
    "cathy": ["catherine"],
    "kate": ["katherine"],
    "kathy": ["katherine", "kathryn"],
    "katie": ["katherine", "kathryn"],
    "beth": ["elizabeth"],
    "liz": ["elizabeth"],
    "lisa": ["elizabeth"],
    "jenny": ["jennifer"],
    "jen": ["jennifer"],
    "sue": ["susan"],
    "susie": ["susan"],
    "meg": ["margaret"],
    "maggie": ["margaret"],
    "peggy": ["margaret"],
    "pam": ["pamela"],
    "sandy": ["sandra"],
    "barb": ["barbara"],
    "deb": ["deborah"],
    "debbie": ["deborah"],
    "chris": ["christine", "christina"],
    "chrissy": ["christine"],
    "tina": ["christina"],
    "nat": ["natalie"],
    "nikki": ["nicole"],
    "ally": ["allison"],
    "ali": ["alison"],
    "sam": ["samantha"],
    "cindy": ["cynthia"],
    "terry": ["theresa"],
    "teri": ["theresa"],
    "trish": ["patricia"],
    "pat": ["patricia"],
    "vicky": ["victoria"],
    "vicki": ["victoria"],
    "becky": ["rebecca"],
    "becca": ["rebecca"],
    "steph": ["stephanie"],
    "cassie": ["cassandra"],
    "ronnie": ["veronica"],
    "mel": ["melanie", "melissa"],
    "jess": ["jessica"],
    "lin": ["linda"],
    "angie": ["angela"],
    "abby": ["abigail"],
    "lexi": ["alexis"],
    "alex": ["alexandra", "alexis"],
    "aly": ["alyssa"],
    "ellie": ["eleanor", "elaine"],
}


# ---------------------------------------------------------------------------
# Styling constants (match usta-team-scout look)
# ---------------------------------------------------------------------------

NAVY = "1F3864"
LIGHT_BLUE = "D6E4F0"
MID_BLUE = "2E75B6"
WIN_GREEN = "E2EFDA"
LOSS_RED = "FCE4D6"
WARNING_YELLOW = "FFF2CC"
WHITE = "FFFFFF"


# ---------------------------------------------------------------------------
# Data classes
# ---------------------------------------------------------------------------


@dataclass
class SearchHit:
    name: str
    href: str           # relative, e.g. "/adult/profile.aspx?playername=X&s=2"
    location: str       # e.g. "CITY, STATE"
    gender: str
    ntrp: str           # e.g. "3.5 C"
    updated: str


@dataclass
class PlayerInfo:
    name: str
    location: str
    ntrp_level: str
    rating_type: str            # "C", "S", or "-"
    dynamic_rating: Optional[float]
    rating_as_of: str
    profile_url: str
    match_history_urls: Dict[int, str] = field(default_factory=dict)  # year -> url


@dataclass
class Opponent:
    name: str
    rating: Optional[float]


@dataclass
class Match:
    date: date
    court: str                  # S1 / S2 / D1 / D2 / D3
    is_singles: bool
    league: str
    level: str
    my_team: str
    opponent_team: str
    result: str                 # "W" or "L"
    score: str                  # e.g. "6-1 6-4"
    partner: Optional[Opponent]
    opponents: List[Opponent]
    match_difficulty: Optional[float]
    dynamic_rating_after: Optional[float]
    rating_trend_hint: str      # "up" / "down" / "flat"


# ---------------------------------------------------------------------------
# HTTP helpers (curl-based to match usta-team-scout style)
# ---------------------------------------------------------------------------


def _curl(args: List[str]) -> str:
    return subprocess.check_output(args, text=True)


def http_get(url: str) -> str:
    safe_url = quote(url, safe=":/?&=%")
    return _curl(["curl", "-L", "--silent", safe_url])


def http_post(url: str, form: Dict[str, str]) -> str:
    data = "&".join(f"{k}={quote(v)}" for k, v in form.items())
    return _curl(["curl", "-L", "--silent", "-X", "POST", "-d", data, url])


def soup(html: str) -> BeautifulSoup:
    return BeautifulSoup(html, "lxml")


# ---------------------------------------------------------------------------
# Search
# ---------------------------------------------------------------------------


def _parse_search_results(html: str) -> List[SearchHit]:
    s = soup(html)
    hits: List[SearchHit] = []
    for table in s.find_all("table"):
        headers = [th.get_text(strip=True) for th in table.find_all("th")]
        if "Player Name" in headers and "Location" in headers:
            for row in table.find_all("tr"):
                cells = row.find_all("td")
                if len(cells) < 4:
                    continue
                a = cells[0].find("a")
                if not a or not a.get("href"):
                    continue
                hits.append(SearchHit(
                    name=a.get_text(strip=True),
                    href=a["href"],
                    location=cells[1].get_text(" ", strip=True),
                    gender=cells[2].get_text(strip=True),
                    ntrp=" ".join(cells[3].get_text(" ", strip=True).split()),
                    updated=cells[4].get_text(strip=True) if len(cells) > 4 else "",
                ))
            break
    return hits


def search_players(first: str, last: str) -> List[SearchHit]:
    html = http_post(SEARCH_URL, {"firstname": first, "lastname": last})
    return _parse_search_results(html)


def search_players_with_fallback(first: str, last: str, location_hint: Optional[str]) -> Tuple[SearchHit, str]:
    """Search with automatic nickname expansion on failure. Returns (hit, resolved_first_name)."""
    # 1. Try exact name
    hits = search_players(first, last)
    if hits:
        try:
            return pick_search_hit(hits, location_hint), first
        except SystemExit:
            pass

    # 2. Try nickname expansions
    expansions = NICKNAME_MAP.get(first.lower(), [])
    for expanded in expansions:
        hits = search_players(expanded.title(), last)
        if hits:
            print(f"[player-scout] '{first}' not found; trying nickname expansion '{expanded.title()}'",
                  file=sys.stderr)
            try:
                return pick_search_hit(hits, location_hint), expanded.title()
            except SystemExit:
                continue

    raise SystemExit(
        f"No players found for '{first} {last}' or any nickname expansions "
        f"({', '.join(e.title() for e in expansions) or 'none known'}). "
        "Check spelling or use --team to search by roster."
    )


def search_players_via_team(first: str, last: str, team_name: str, year: int) -> Tuple[SearchHit, str]:
    """Find a player by fuzzy-matching their name against a team's roster (2.2)."""
    from urllib.parse import quote as _quote
    team_url = f"{BASE_URL}/adult/teamprofile.aspx?year={year}&teamname={_quote(team_name, safe='')}"
    html = http_get(team_url)
    if "No team found" in html or html.count("<table") < 3:
        # Try &s=1, &s=2, &s=3
        for s_val in [1, 2, 3]:
            html = http_get(team_url + f"&s={s_val}")
            if html.count("<table") >= 3:
                break

    s = soup(html)
    target = f"{first} {last}".lower()

    best_hit: Optional[SearchHit] = None
    best_score = 0
    for a in s.find_all("a", href=lambda h: h and "/adult/profile.aspx?playername=" in h):
        roster_name = a.get_text(strip=True)
        # Simple token overlap score
        roster_tokens = set(roster_name.lower().split())
        target_tokens = set(target.split())
        score = len(roster_tokens & target_tokens)
        if score > best_score:
            best_score = score
            best_hit = SearchHit(
                name=roster_name,
                href=a["href"],
                location="",
                gender="",
                ntrp="",
                updated="",
            )

    if best_hit is None or best_score == 0:
        raise SystemExit(
            f"Could not find '{first} {last}' on team '{team_name}'. "
            "Check team name spelling."
        )

    print(f"[player-scout] Matched '{first} {last}' → '{best_hit.name}' via team roster",
          file=sys.stderr)
    return best_hit, best_hit.name.split()[0]


def pick_search_hit(hits: List[SearchHit], location_hint: Optional[str]) -> SearchHit:
    if not hits:
        raise SystemExit("No players found for that name. Check spelling.")
    if len(hits) == 1:
        return hits[0]
    if location_hint:
        key = location_hint.lower().strip()
        matches = [h for h in hits if key in h.location.lower()]
        if len(matches) == 1:
            return matches[0]
        if len(matches) > 1:
            msg = "\n".join(f"  - {h.name} ({h.location}, {h.ntrp})" for h in matches)
            raise SystemExit(
                f"Multiple players match '{location_hint}':\n{msg}\n"
                "Please pass a more specific --location."
            )
    msg = "\n".join(f"  - {h.name} ({h.location}, {h.ntrp})" for h in hits)
    raise SystemExit(
        "Multiple players found — please pass --location or --s to disambiguate:\n"
        f"{msg}"
    )


# ---------------------------------------------------------------------------
# Profile
# ---------------------------------------------------------------------------


_DR_RE = re.compile(r"(\d\.\d{2,4})")


def parse_profile(profile_url: str) -> PlayerInfo:
    html = http_get(profile_url)
    s = soup(html)

    name_anchor = s.find("a", href=re.compile(r"/adult/profile\.aspx"))
    name = name_anchor.get_text(strip=True) if name_anchor else ""
    loc_match = (re.search(r"\(([^)]+)\)", name_anchor.parent.get_text(" ", strip=True))
                 if name_anchor else None)
    location = loc_match.group(1).strip() if loc_match else ""

    text = s.get_text(" ", strip=True)

    ntrp_level = ""
    rating_type = "-"
    ntrp_m = re.search(r"(\d\.\d)\s*([CSA])\b", text)
    if ntrp_m:
        ntrp_level, rating_type = ntrp_m.group(1), ntrp_m.group(2)

    dr: Optional[float] = None
    dr_m = re.search(r"Estimated Dynamic Rating.*?(\d\.\d{3,4})", text)
    if dr_m:
        dr = float(dr_m.group(1))

    as_of_m = re.search(r"as of\s+(\d{1,2}/\d{1,2}/\d{4})", text, re.IGNORECASE)
    if as_of_m:
        rating_as_of = as_of_m.group(1)
    else:
        near_ntrp = re.search(r"\d\.\d\s*[CSA]\s+(\d{1,2}/\d{1,2}/\d{4})", text)
        rating_as_of = near_ntrp.group(1) if near_ntrp else ""

    year_urls: Dict[int, str] = {}
    for a in s.find_all("a", href=re.compile(r"/adult/matchhistory\.aspx\?year=\d{4}")):
        href = a["href"]
        ym = re.search(r"year=(\d{4})", href)
        if ym:
            year_urls[int(ym.group(1))] = BASE_URL + href.replace(" ", "%20")

    return PlayerInfo(
        name=name,
        location=location,
        ntrp_level=ntrp_level,
        rating_type=rating_type,
        dynamic_rating=dr,
        rating_as_of=rating_as_of,
        profile_url=profile_url,
        match_history_urls=year_urls,
    )


# ---------------------------------------------------------------------------
# Match history parsing
# ---------------------------------------------------------------------------


def _float_or_none(s: str) -> Optional[float]:
    s = s.strip()
    if not s or s.startswith("-") or s == "—":
        return None
    m = _DR_RE.search(s)
    return float(m.group(1)) if m else None


def _parse_opponents(cell) -> List[Opponent]:
    out: List[Opponent] = []
    html = cell.decode_contents()
    parts = re.split(r"<br\s*/?>", html, flags=re.IGNORECASE)
    for part in parts:
        sub = BeautifulSoup(part, "lxml")
        a = sub.find("a")
        if not a:
            continue
        name = a.get_text(strip=True)
        text = sub.get_text(" ", strip=True)
        rating = _float_or_none(text.replace(name, "", 1))
        out.append(Opponent(name=name, rating=rating))
    if not out:
        for a in cell.find_all("a"):
            out.append(Opponent(name=a.get_text(strip=True), rating=None))
    return out


def parse_match_history(html: str) -> List[Match]:
    s = soup(html)
    matches: List[Match] = []

    for div in s.select("div.container496"):
        tbl = div.find("table")
        if not tbl:
            continue
        rows = tbl.find_all("tr", recursive=False) or tbl.find_all("tr")
        if len(rows) < 4:
            continue

        row0 = rows[0].find_all(["th", "td"])
        if len(row0) < 3:
            continue
        date_text = row0[0].get_text(strip=True)
        court = row0[1].get_text(strip=True)
        league_blob = row0[2].get_text(" ", strip=True)
        try:
            m_date = datetime.strptime(date_text, "%m/%d/%Y").date()
        except ValueError:
            continue
        league_parts = league_blob.split()
        level = league_parts[-1] if league_parts else ""
        league = " ".join(league_parts[:-1]) if len(league_parts) > 1 else league_blob

        row1 = rows[1].find_all(["th", "td"])
        if len(row1) < 3:
            continue
        my_team = row1[0].get_text(" ", strip=True).split("Pacific NW")[0].strip()
        my_team = re.sub(r"\s+", " ", my_team).strip()
        result_raw = row1[1].get_text(strip=True).upper()
        result = ("W" if result_raw.startswith("W") else
                  "L" if result_raw.startswith("L") else result_raw)
        opp_team = row1[2].get_text(" ", strip=True)
        opp_team = re.sub(r"\s+Pacific NW.*$", "", opp_team).strip()

        row2 = rows[2].find_all(["th", "td"])
        if len(row2) < 3:
            continue
        partner_cell, score_cell, opp_cell = row2[0], row2[1], row2[2]
        partners = _parse_opponents(partner_cell)
        partner = partners[0] if partners else None
        score = " ".join(line.strip() for line in score_cell.get_text("\n").splitlines()
                         if line.strip())
        opponents = _parse_opponents(opp_cell)

        row3_text = rows[3].get_text(" ", strip=True)
        diff_m = re.search(r"Match:\s*([\d.]+|-+)", row3_text)
        rate_m = re.search(r"Rating:\s*([\d.]+|-+)", row3_text)
        match_diff = _float_or_none(diff_m.group(1)) if diff_m else None
        post_dr = _float_or_none(rate_m.group(1)) if rate_m else None

        row3_html = str(rows[3])
        if "#DD0000" in row3_html:
            trend = "down"
        elif "#00DD00" in row3_html:
            trend = "up"
        else:
            trend = "flat"

        is_singles = court.upper().startswith("S")
        matches.append(Match(
            date=m_date,
            court=court,
            is_singles=is_singles,
            league=league.strip(),
            level=level,
            my_team=my_team,
            opponent_team=opp_team,
            result=result,
            score=score,
            partner=partner,
            opponents=opponents,
            match_difficulty=match_diff,
            dynamic_rating_after=post_dr,
            rating_trend_hint=trend,
        ))

    return matches


def dedupe_matches(matches: List[Match]) -> List[Match]:
    seen = {}
    for m in matches:
        key = (m.date, m.court, tuple(sorted(o.name for o in m.opponents)))
        seen[key] = m
    return sorted(seen.values(), key=lambda m: m.date, reverse=True)


# ---------------------------------------------------------------------------
# Analysis
# ---------------------------------------------------------------------------


def _mean(xs: List[float]) -> Optional[float]:
    xs = [x for x in xs if x is not None]
    return sum(xs) / len(xs) if xs else None


def _median(xs: List[float]) -> Optional[float]:
    xs = sorted(x for x in xs if x is not None)
    if not xs:
        return None
    mid = len(xs) // 2
    if len(xs) % 2:
        return xs[mid]
    return (xs[mid - 1] + xs[mid]) / 2


def _fmt(x: Optional[float], digits: int = 2) -> str:
    return f"{x:.{digits}f}" if x is not None else "—"


def _slope(pairs: List[Tuple[date, float]]) -> Optional[float]:
    pts = [(p[0].toordinal(), p[1]) for p in pairs if p[1] is not None]
    if len(pts) < 3:
        return None
    n = len(pts)
    mx = sum(p[0] for p in pts) / n
    my = sum(p[1] for p in pts) / n
    num = sum((p[0] - mx) * (p[1] - my) for p in pts)
    den = sum((p[0] - mx) ** 2 for p in pts)
    return num / den if den else None


def analyze(matches: List[Match], player_dr: Optional[float]) -> Dict:
    total = len(matches)
    singles = [m for m in matches if m.is_singles]
    doubles = [m for m in matches if not m.is_singles]
    wins = [m for m in matches if m.result == "W"]
    losses = [m for m in matches if m.result == "L"]

    last10 = list(reversed(matches[:10]))
    form = "".join(m.result for m in last10)

    streak_count = 0
    streak_type = None
    for m in matches:
        if streak_type is None:
            streak_type = m.result
            streak_count = 1
        elif m.result == streak_type:
            streak_count += 1
        else:
            break
    streak = f"{streak_type}{streak_count}" if streak_type else "—"

    diffs = [m.match_difficulty for m in matches if m.match_difficulty is not None]
    buckets: Dict[str, int] = {}
    for d in diffs:
        key = f"{(int(d * 4) / 4):.2f}"
        buckets[key] = buckets.get(key, 0) + 1

    def record_vs(pred) -> Tuple[int, int]:
        w = sum(1 for m in matches
                if m.match_difficulty is not None
                and pred(m.match_difficulty) and m.result == "W")
        l = sum(1 for m in matches
                if m.match_difficulty is not None
                and pred(m.match_difficulty) and m.result == "L")
        return w, l

    if player_dr is not None:
        vs_higher = record_vs(lambda d: d >= player_dr + 0.25)
        vs_equal = record_vs(lambda d: abs(d - player_dr) < 0.25)
        vs_lower = record_vs(lambda d: d <= player_dr - 0.25)
    else:
        vs_higher = vs_equal = vs_lower = (0, 0)

    dr_series = [(m.date, m.dynamic_rating_after) for m in reversed(matches)
                 if m.dynamic_rating_after is not None]
    slope_per_day = _slope(dr_series) or 0.0
    if slope_per_day > 0.0007:
        trend = "rising"
    elif slope_per_day < -0.0007:
        trend = "falling"
    else:
        trend = "steady"

    dr_values = [v for _, v in dr_series]
    dr_first = dr_series[0][1] if dr_series else None
    dr_last = dr_series[-1][1] if dr_series else None
    dr_max = max(dr_values) if dr_values else None
    dr_min = min(dr_values) if dr_values else None

    court_counts: Dict[str, int] = {}
    for m in matches:
        court_counts[m.court] = court_counts.get(m.court, 0) + 1

    partner_stats: Dict[str, Dict] = {}
    for m in doubles:
        if not m.partner:
            continue
        ps = partner_stats.setdefault(m.partner.name, {
            "matches": 0, "wins": 0, "ratings": []
        })
        ps["matches"] += 1
        if m.result == "W":
            ps["wins"] += 1
        if m.partner.rating is not None:
            ps["ratings"].append(m.partner.rating)
    partners_sorted = sorted(
        partner_stats.items(), key=lambda kv: (-kv[1]["matches"], kv[0])
    )

    top_wins = sorted(
        [m for m in wins if m.match_difficulty is not None],
        key=lambda m: m.match_difficulty, reverse=True
    )[:3]
    upset_losses = sorted(
        [m for m in losses if m.match_difficulty is not None],
        key=lambda m: m.match_difficulty
    )[:3]

    months = sorted({(m.date.year, m.date.month) for m in matches})
    if len(matches) >= 2:
        gaps = [(matches[i - 1].date - matches[i].date).days
                for i in range(1, len(matches))]
        longest_gap = max(gaps)
    else:
        longest_gap = 0

    level_counts: Dict[str, int] = {}
    for m in matches:
        key = f"{m.league.strip()} {m.level}".strip() or "Unknown"
        level_counts[key] = level_counts.get(key, 0) + 1

    # Level-switching detection (2.3): detect if player is competing at a lower
    # level this year vs prior year
    current_year = date.today().year
    ntrp_re = re.compile(r"(\d\.\d)")
    def _extract_ntrp(level_str: str) -> Optional[float]:
        m = ntrp_re.search(level_str)
        return float(m.group(1)) if m else None

    prior_levels = set()
    current_levels = set()
    for m in matches:
        ntrp = _extract_ntrp(m.level)
        if ntrp is None:
            continue
        if m.date.year < current_year:
            prior_levels.add(ntrp)
        else:
            current_levels.add(ntrp)

    stepping_down: Optional[str] = None
    if prior_levels and current_levels:
        prior_max = max(prior_levels)
        current_max = max(current_levels)
        if current_max < prior_max:
            prior_record_at_higher = [(m.result, m.level) for m in matches
                                       if _extract_ntrp(m.level) == prior_max
                                       and m.date.year < current_year]
            prior_w = sum(1 for r, _ in prior_record_at_higher if r == "W")
            prior_l = len(prior_record_at_higher) - prior_w
            stepping_down = (
                f"Playing {current_max} in {current_year} after competing at "
                f"{prior_max} in prior season(s). Prior record at {prior_max}: "
                f"{prior_w}W-{prior_l}L. DR may look strong for this level — "
                f"verify actual competitive history."
            )

    return {
        "total": total,
        "singles_total": len(singles),
        "doubles_total": len(doubles),
        "singles_wl": (sum(1 for m in singles if m.result == "W"),
                       sum(1 for m in singles if m.result == "L")),
        "doubles_wl": (sum(1 for m in doubles if m.result == "W"),
                       sum(1 for m in doubles if m.result == "L")),
        "wins": len(wins),
        "losses": len(losses),
        "form": form,
        "streak": streak,
        "diff_mean": _mean(diffs),
        "diff_median": _median(diffs),
        "diff_min": min(diffs) if diffs else None,
        "diff_max": max(diffs) if diffs else None,
        "diff_buckets": dict(sorted(buckets.items())),
        "vs_higher": vs_higher,
        "vs_equal": vs_equal,
        "vs_lower": vs_lower,
        "dr_first": dr_first,
        "dr_last": dr_last,
        "dr_max": dr_max,
        "dr_min": dr_min,
        "dr_slope_per_day": slope_per_day,
        "trend": trend,
        "dr_series": dr_series,
        "court_counts": court_counts,
        "partners": partners_sorted,
        "top_wins": top_wins,
        "upset_losses": upset_losses,
        "months_active": len(months),
        "longest_gap": longest_gap,
        "level_counts": level_counts,
        "stepping_down": stepping_down,
    }


# ---------------------------------------------------------------------------
# docx helpers
# ---------------------------------------------------------------------------


def _shade_cell(cell, color_hex: str) -> None:
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tcPr.append(shd)


def _set_cell_text(cell, text: str, *, bold: bool = False,
                   color_hex: Optional[str] = None,
                   align=WD_ALIGN_PARAGRAPH.LEFT,
                   font_size: Optional[int] = None) -> None:
    cell.text = ""
    para = cell.paragraphs[0]
    para.alignment = align
    run = para.add_run(text)
    run.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if font_size is not None:
        run.font.size = Pt(font_size)


def _pct(w: int, l: int) -> str:
    n = w + l
    return f"{100 * w / n:.1f}%" if n else "—"


def _opponents_str(ops: List[Opponent]) -> str:
    parts = []
    for o in ops:
        r = f"{o.rating:.2f}" if o.rating is not None else "—"
        parts.append(f"{o.name} ({r})")
    return " / ".join(parts)


def _opponent_rating_avg(ops: List[Opponent]) -> Optional[float]:
    rs = [o.rating for o in ops if o.rating is not None]
    return sum(rs) / len(rs) if rs else None


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    sizes = {0: 22, 1: 16, 2: 13, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = RGBColor.from_string(NAVY)
    if level <= 1:
        p_fmt = p.paragraph_format
        p_fmt.space_before = Pt(12)
        p_fmt.space_after = Pt(6)


def _add_kv_table(doc: Document, rows: List[Tuple[str, str]]) -> None:
    t = doc.add_table(rows=len(rows), cols=2)
    t.alignment = WD_TABLE_ALIGNMENT.LEFT
    t.autofit = True
    for i, (k, v) in enumerate(rows):
        _set_cell_text(t.rows[i].cells[0], k, bold=True, font_size=10)
        _set_cell_text(t.rows[i].cells[1], v, font_size=10)
        _shade_cell(t.rows[i].cells[0], LIGHT_BLUE)


def _add_header_row(table, headers: List[str]) -> None:
    for cell, h in zip(table.rows[0].cells, headers):
        _set_cell_text(cell, h, bold=True, color_hex=WHITE, font_size=10,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _shade_cell(cell, NAVY)


# ---------------------------------------------------------------------------
# Report rendering (.docx)
# ---------------------------------------------------------------------------


def render_docx(player: PlayerInfo, matches: List[Match], analysis: Dict,
                lookback_start: date, lookback_end: date,
                out_path: Path) -> None:
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    # Default font tweak
    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    # ---- Title block -----------------------------------------------------
    title = doc.add_paragraph()
    run = title.add_run(f"Player Scout: {player.name}")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    header_bits = []
    if player.location:
        header_bits.append(f"Location: {player.location}")
    if player.ntrp_level:
        header_bits.append(f"NTRP: {player.ntrp_level} {player.rating_type}".strip())
    if player.dynamic_rating is not None:
        dr_text = f"Dynamic Rating: {_fmt(player.dynamic_rating, 4)}"
        if player.rating_as_of:
            dr_text += f" (as of {player.rating_as_of})"
        header_bits.append(dr_text)
    sub = doc.add_paragraph(" · ".join(header_bits))
    sub.runs[0].font.size = Pt(11)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Lookback: {lookback_start.isoformat()} → {lookback_end.isoformat()} "
        f"({(lookback_end - lookback_start).days // 30} months)"
    ).font.size = Pt(9)
    meta.add_run("\n").font.size = Pt(9)
    src_run = meta.add_run(f"Source: {player.profile_url}")
    src_run.font.size = Pt(9)
    src_run.font.color.rgb = RGBColor.from_string(MID_BLUE)

    if player.rating_type == "S":
        warn = doc.add_paragraph()
        w_run = warn.add_run(
            "⚠ Self-rated — may play above listed NTRP. Treat with caution."
        )
        w_run.bold = True
        w_run.font.color.rgb = RGBColor.from_string("C00000")

    # ---- TL;DR -----------------------------------------------------------
    _add_heading(doc, "TL;DR", level=1)
    tldr: List[str] = [
        f"{analysis['total']} matches in window "
        f"({analysis['singles_total']} singles, {analysis['doubles_total']} doubles), "
        f"{analysis['wins']}–{analysis['losses']} overall."
    ]
    if analysis["diff_mean"] is not None:
        tldr.append(
            f"Typical opponent strength ~{_fmt(analysis['diff_mean'])} "
            f"(range {_fmt(analysis['diff_min'])}–{_fmt(analysis['diff_max'])})."
        )
    if analysis["dr_first"] is not None and analysis["dr_last"] is not None:
        delta = analysis["dr_last"] - analysis["dr_first"]
        sign = "+" if delta >= 0 else ""
        tldr.append(
            f"Dynamic rating {analysis['trend']}: "
            f"{_fmt(analysis['dr_first'], 3)} → {_fmt(analysis['dr_last'], 3)} "
            f"({sign}{delta:.3f})."
        )
    if analysis["form"]:
        tldr.append(
            f"Last {len(analysis['form'])} matches: {analysis['form']} "
            f"(current streak {analysis['streak']})."
        )
    if analysis["vs_higher"] != (0, 0):
        w, l = analysis["vs_higher"]
        tldr.append(f"Vs opponents ≥0.25 above their DR: {w}–{l}.")
    for b in tldr:
        doc.add_paragraph(b, style="List Bullet")
    if analysis.get("stepping_down"):
        warn = doc.add_paragraph()
        wr = warn.add_run(f"⚠ STEPPING DOWN: {analysis['stepping_down']}")
        wr.bold = True
        wr.font.color.rgb = RGBColor.from_string("C00000")

    # ---- Volume & Activity ----------------------------------------------
    _add_heading(doc, "Volume & Activity", level=1)
    _add_kv_table(doc, [
        ("Total matches", str(analysis["total"])),
        ("Singles / Doubles",
         f"{analysis['singles_total']} / {analysis['doubles_total']}"),
        ("Active months", str(analysis["months_active"])),
        ("Longest gap between matches", f"{analysis['longest_gap']} days"),
    ])
    doc.add_paragraph()
    mix = doc.add_paragraph()
    mix.add_run("Play mix by league/level:").bold = True
    for k, v in sorted(analysis["level_counts"].items(), key=lambda kv: -kv[1]):
        doc.add_paragraph(f"{k}: {v}", style="List Bullet")

    # ---- Results --------------------------------------------------------
    _add_heading(doc, "Results", level=1)
    res_table = doc.add_table(rows=4, cols=4)
    res_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _add_header_row(res_table, ["Split", "W", "L", "Win %"])
    sw, sl = analysis["singles_wl"]
    dw, dl = analysis["doubles_wl"]
    for i, (label, w, l) in enumerate([
        ("Overall", analysis["wins"], analysis["losses"]),
        ("Singles", sw, sl),
        ("Doubles", dw, dl),
    ], start=1):
        cells = res_table.rows[i].cells
        _set_cell_text(cells[0], label, font_size=10)
        _set_cell_text(cells[1], str(w), font_size=10,
                       align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(cells[2], str(l), font_size=10,
                       align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(cells[3], _pct(w, l), font_size=10,
                       align=WD_ALIGN_PARAGRAPH.RIGHT)

    doc.add_paragraph()
    form_p = doc.add_paragraph()
    form_p.add_run(f"Last {len(analysis['form'])} matches: ").bold = True
    form_p.add_run(analysis["form"] or "—")
    streak_p = doc.add_paragraph()
    streak_p.add_run("Current streak: ").bold = True
    streak_p.add_run(analysis["streak"])

    # ---- Opponent Strength ---------------------------------------------
    _add_heading(doc, "Opponent Strength", level=1)
    _add_kv_table(doc, [
        ("Mean match difficulty", _fmt(analysis["diff_mean"])),
        ("Median", _fmt(analysis["diff_median"])),
        ("Range", f"{_fmt(analysis['diff_min'])} → {_fmt(analysis['diff_max'])}"),
    ])
    if player.dynamic_rating is not None:
        doc.add_paragraph()
        rel_p = doc.add_paragraph()
        rel_p.add_run(
            f"Relative to current DR ({_fmt(player.dynamic_rating, 3)}):"
        ).bold = True
        rel_table = doc.add_table(rows=4, cols=4)
        _add_header_row(rel_table, ["Bucket", "W", "L", "Win %"])
        for i, (label, (w, l)) in enumerate([
            ("Tougher (opp ≥ DR+0.25)", analysis["vs_higher"]),
            ("Peer (within ±0.25)", analysis["vs_equal"]),
            ("Easier (opp ≤ DR−0.25)", analysis["vs_lower"]),
        ], start=1):
            cells = rel_table.rows[i].cells
            _set_cell_text(cells[0], label, font_size=10)
            _set_cell_text(cells[1], str(w), font_size=10,
                           align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(cells[2], str(l), font_size=10,
                           align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(cells[3], _pct(w, l), font_size=10,
                           align=WD_ALIGN_PARAGRAPH.RIGHT)

    if analysis["diff_buckets"]:
        doc.add_paragraph()
        bp = doc.add_paragraph()
        bp.add_run("Opponent difficulty distribution (0.25 buckets):").bold = True
        bucket_table = doc.add_table(rows=1 + len(analysis["diff_buckets"]), cols=2)
        _add_header_row(bucket_table, ["Bucket", "Count"])
        for i, (k, v) in enumerate(analysis["diff_buckets"].items(), start=1):
            cells = bucket_table.rows[i].cells
            _set_cell_text(cells[0], k, font_size=10)
            _set_cell_text(cells[1], str(v), font_size=10,
                           align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ---- Rating Trend ---------------------------------------------------
    _add_heading(doc, "Rating Trend", level=1)
    # One-line trend summary
    if analysis["dr_first"] is not None and analysis["dr_last"] is not None:
        n_dr = len(analysis["dr_series"])
        delta = analysis["dr_last"] - analysis["dr_first"]
        sign = "+" if delta >= 0 else ""
        trend_summary = doc.add_paragraph()
        ts_run = trend_summary.add_run(
            f"DR Trend: {analysis['dr_first']:.2f} → {analysis['dr_last']:.2f} "
            f"({sign}{delta:.2f} over last {n_dr} matches, {analysis['trend']})"
        )
        ts_run.bold = True
    _add_kv_table(doc, [
        ("First DR in window", _fmt(analysis["dr_first"], 4)),
        ("Last DR in window", _fmt(analysis["dr_last"], 4)),
        ("Max / Min",
         f"{_fmt(analysis['dr_max'], 4)} / {_fmt(analysis['dr_min'], 4)}"),
        ("Trend",
         f"{analysis['trend']} (slope {analysis['dr_slope_per_day']:+.5f}/day)"),
    ])

    # ---- Court & Partners ----------------------------------------------
    _add_heading(doc, "Court & Partners", level=1)
    if analysis["court_counts"]:
        cc_p = doc.add_paragraph()
        cc_p.add_run("Court position counts:").bold = True
        court_table = doc.add_table(rows=1 + len(analysis["court_counts"]), cols=2)
        _add_header_row(court_table, ["Court", "Matches"])
        for i, (court, n) in enumerate(sorted(analysis["court_counts"].items()),
                                       start=1):
            cells = court_table.rows[i].cells
            _set_cell_text(cells[0], court, font_size=10)
            _set_cell_text(cells[1], str(n), font_size=10,
                           align=WD_ALIGN_PARAGRAPH.RIGHT)

    if analysis["partners"]:
        doc.add_paragraph()
        pp = doc.add_paragraph()
        pp.add_run("Most frequent doubles partners:").bold = True
        top_partners = analysis["partners"][:5]
        p_table = doc.add_table(rows=1 + len(top_partners), cols=5)
        _add_header_row(p_table, ["Partner", "Matches", "Wins", "Win %",
                                  "Avg Rating"])
        for i, (name, ps) in enumerate(top_partners, start=1):
            avg = (sum(ps["ratings"]) / len(ps["ratings"])
                   if ps["ratings"] else None)
            n = ps["matches"]
            w = ps["wins"]
            cells = p_table.rows[i].cells
            _set_cell_text(cells[0], name, font_size=10)
            _set_cell_text(cells[1], str(n), font_size=10,
                           align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(cells[2], str(w), font_size=10,
                           align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(cells[3], _pct(w, n - w), font_size=10,
                           align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(cells[4], _fmt(avg), font_size=10,
                           align=WD_ALIGN_PARAGRAPH.RIGHT)

    # ---- Notable Matches -----------------------------------------------
    if analysis["top_wins"] or analysis["upset_losses"]:
        _add_heading(doc, "Notable Matches", level=1)
        if analysis["top_wins"]:
            tw = doc.add_paragraph()
            tw.add_run("Biggest wins (by opponent strength):").bold = True
            for m in analysis["top_wins"]:
                doc.add_paragraph(
                    f"{m.date.isoformat()} · {m.court} · vs "
                    f"{_opponents_str(m.opponents)} · score {m.score} · "
                    f"match diff {_fmt(m.match_difficulty)}",
                    style="List Bullet",
                )
        if analysis["upset_losses"]:
            ul = doc.add_paragraph()
            ul.add_run("Losses to notably weaker opponents:").bold = True
            for m in analysis["upset_losses"]:
                doc.add_paragraph(
                    f"{m.date.isoformat()} · {m.court} · vs "
                    f"{_opponents_str(m.opponents)} · score {m.score} · "
                    f"match diff {_fmt(m.match_difficulty)}",
                    style="List Bullet",
                )

    # ---- Full Match Log -------------------------------------------------
    _add_heading(doc, "Full Match Log", level=1)
    headers = ["Date", "Type", "Court", "Opp Team", "Opponent(s)",
               "Opp Rating", "Score", "W/L", "Post DR", "DR Δ"]
    log_table = doc.add_table(rows=1 + len(matches), cols=len(headers))
    _add_header_row(log_table, headers)
    # Build ordered DR series for delta calculation (matches are newest-first)
    ordered_matches = list(reversed(matches))
    prev_dr_map: Dict[int, Optional[float]] = {}
    for idx, m in enumerate(ordered_matches):
        prev_dr_map[idx] = ordered_matches[idx - 1].dynamic_rating_after if idx > 0 else None

    for i, m in enumerate(matches, start=1):
        # Map back to ordered index for delta
        ordered_idx = len(matches) - i
        prev_dr = prev_dr_map.get(ordered_idx)
        cur_dr = m.dynamic_rating_after
        if cur_dr is not None and prev_dr is not None:
            delta_val = cur_dr - prev_dr
            dr_delta_str = f"{delta_val:+.2f}"
        else:
            dr_delta_str = "—"

        tpe = "S" if m.is_singles else "D"
        avg = _opponent_rating_avg(m.opponents)
        cells = log_table.rows[i].cells
        _set_cell_text(cells[0], m.date.isoformat(), font_size=9)
        _set_cell_text(cells[1], tpe, font_size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[2], m.court, font_size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[3], m.opponent_team, font_size=9)
        _set_cell_text(cells[4], _opponents_str(m.opponents), font_size=9)
        _set_cell_text(cells[5], _fmt(avg), font_size=9,
                       align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(cells[6], m.score, font_size=9)
        _set_cell_text(cells[7], m.result, font_size=9,
                       align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        _set_cell_text(cells[8], _fmt(cur_dr, 4), font_size=9,
                       align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(cells[9], dr_delta_str, font_size=9,
                       align=WD_ALIGN_PARAGRAPH.RIGHT)
        # Row shading by result
        shade = WIN_GREEN if m.result == "W" else LOSS_RED if m.result == "L" else None
        if shade:
            _shade_cell(cells[7], shade)

    # ---- Footer note ---------------------------------------------------
    doc.add_paragraph()
    footer = doc.add_paragraph()
    fr = footer.add_run(
        f"Scouted {datetime.now().strftime('%Y-%m-%d %H:%M')} from public "
        f"tennisrecord.com data."
    )
    fr.italic = True
    fr.font.size = Pt(9)
    fr.font.color.rgb = RGBColor.from_string("7A7A7A")

    doc.save(str(out_path))


# ---------------------------------------------------------------------------
# CLI
# ---------------------------------------------------------------------------


def print_inline_summary(player: PlayerInfo, matches: List[Match], analysis: Dict) -> None:
    """Print a 5-8 line quick-look summary to stdout (--summary mode, 2.4)."""
    dr_str = f"{player.dynamic_rating:.2f}" if player.dynamic_rating is not None else "—"
    rt = player.rating_type if player.rating_type != "-" else "?"
    print(f"  Player:  {player.name}  ({player.location})")
    print(f"  Rating:  {player.ntrp_level} {rt}  |  DR: {dr_str}"
          + (f"  as of {player.rating_as_of}" if player.rating_as_of else ""))
    if analysis["total"] == 0:
        print("  Matches: No match history on record.")
        if player.rating_type == "S":
            print("  ⚠ Self-rated with no history — true unknown, could play at any level.")
        return
    wl = f"{analysis['wins']}W-{analysis['losses']}L"
    form = analysis["form"] or "—"
    print(f"  Record:  {wl}  ({analysis['total']} matches, {analysis['singles_total']}S/{analysis['doubles_total']}D)")
    print(f"  Form:    {form}  (streak: {analysis['streak']})")
    trend = analysis["trend"]
    if analysis["dr_first"] is not None and analysis["dr_last"] is not None:
        delta = analysis["dr_last"] - analysis["dr_first"]
        sign = "+" if delta >= 0 else ""
        print(f"  DR Trend: {trend} ({analysis['dr_first']:.3f} → {analysis['dr_last']:.3f}, {sign}{delta:.3f})")
    courts = ", ".join(f"{c}×{n}" for c, n in sorted(analysis["court_counts"].items()))
    print(f"  Courts:  {courts}")
    if analysis.get("stepping_down"):
        print(f"  ⚠ STEPPING DOWN: {analysis['stepping_down']}")


def _slugify(s: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", s).strip("_") or "player"


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--first", required=True, help="Player first name")
    parser.add_argument("--last", required=True, help="Player last name")
    parser.add_argument("--location", default=None,
                        help='Location substring, e.g. "CITY, STATE"')
    parser.add_argument("--team", default=None,
                        help="Team name to scope lookup — searches roster for the player (2.2)")
    parser.add_argument("--s", type=int, default=None,
                        help="Optional &s= disambiguator if already known")
    parser.add_argument("--months", type=int, default=24,
                        help="Lookback window in months (default: 24)")
    parser.add_argument("--output", default=None,
                        help="Output filename (placed under reports/)")
    parser.add_argument("--workspace", default=None,
                        help="Workspace root (defaults to repo root)")
    parser.add_argument("--summary", action="store_true",
                        help="Print a quick inline summary instead of generating a full .docx (2.4)")
    args = parser.parse_args(argv)

    if args.workspace:
        workspace = Path(args.workspace).resolve()
    else:
        workspace = Path(__file__).resolve().parents[2]
    reports_dir = workspace / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    if args.s is not None:
        name_qs = quote(f"{args.first} {args.last}")
        profile_url = f"{BASE_URL}/adult/profile.aspx?playername={name_qs}&s={args.s}"
    elif args.team:
        year = date.today().year
        pick, _ = search_players_via_team(args.first, args.last, args.team, year)
        profile_url = BASE_URL + pick.href.replace(" ", "%20")
    else:
        pick, resolved_first = search_players_with_fallback(args.first, args.last, args.location)
        profile_url = BASE_URL + pick.href.replace(" ", "%20")
        if resolved_first != args.first:
            print(f"[player-scout] Resolved '{args.first}' → '{resolved_first}'", file=sys.stderr)

    player = parse_profile(profile_url)

    today = date.today()
    cutoff = today - timedelta(days=args.months * 30)
    needed_years = sorted({cutoff.year + i for i in range(0, today.year - cutoff.year + 1)})

    all_matches: List[Match] = []
    for year in needed_years:
        url = player.match_history_urls.get(year)
        if not url:
            continue
        html = http_get(url)
        all_matches.extend(parse_match_history(html))

    matches = [m for m in dedupe_matches(all_matches)
               if cutoff <= m.date <= today]

    analysis = analyze(matches, player.dynamic_rating)

    # 2.5 — Self-rated zero-history: flag and skip report
    is_true_unknown = (
        player.rating_type == "S"
        and analysis["total"] == 0
        and player.dynamic_rating is None
    )
    if is_true_unknown:
        print(f"\n⚠  TRUE UNKNOWN: {player.name}")
        print(f"   Self-rated ({player.ntrp_level} S) with no USTA match history on record.")
        print(f"   Cannot assess actual level. Treat as a wildcard — could play well above or below {player.ntrp_level}.")
        if player.rating_as_of:
            print(f"   Rating assigned: {player.rating_as_of}")
        print("\n   No report generated.")
        return 0

    # 2.4 — Inline summary mode: print and exit without generating docx
    if args.summary:
        print_inline_summary(player, matches, analysis)
        return 0

    if args.output:
        fname = Path(args.output).name
        if not fname.lower().endswith(".docx"):
            fname = fname + ".docx"
    else:
        loc_slug = _slugify(player.location.split(",")[0]) if player.location else "loc"
        fname = (f"player_{_slugify(args.first)}_{_slugify(args.last)}_"
                 f"{loc_slug}_{today.strftime('%Y%m%d')}.docx")
    out_path = reports_dir / fname

    render_docx(player, matches, analysis, cutoff, today, out_path)

    print(f"[player-scout] Wrote {out_path}")
    print(f"[player-scout] {analysis['total']} matches in {args.months}m window "
          f"({analysis['singles_total']}S / {analysis['doubles_total']}D); "
          f"record {analysis['wins']}-{analysis['losses']}")
    return 0


if __name__ == "__main__":
    sys.exit(main())
