"""Render a single-player USTA scouting report from pre-fetched JSON data.

Usage:
    python3 skills/usta-player-scout/scripts/player_report.py --input player_data.json

The input JSON is prepared outside Python from fetch_webpage results. This
script performs no outbound network requests and writes a Word report under the
workspace-local reports/ folder.
"""

from __future__ import annotations

import argparse
import json
import re
import sys
from dataclasses import dataclass, field
from datetime import date, datetime, timedelta
from pathlib import Path
from typing import Dict, List, Optional, Tuple

from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


NAVY = "1F3864"
LIGHT_BLUE = "D6E4F0"
MID_BLUE = "2E75B6"
WIN_GREEN = "E2EFDA"
LOSS_RED = "FCE4D6"
WHITE = "FFFFFF"


@dataclass
class PlayerInfo:
    name: str
    location: str
    ntrp_level: str
    rating_type: str
    dynamic_rating: Optional[float]
    rating_as_of: str
    profile_url: str
    match_history_urls: Dict[int, str] = field(default_factory=dict)


@dataclass
class Opponent:
    name: str
    rating: Optional[float]


@dataclass
class Match:
    date: date
    court: str
    is_singles: bool
    league: str
    level: str
    my_team: str
    opponent_team: str
    result: str
    score: str
    partner: Optional[Opponent]
    opponents: List[Opponent]
    match_difficulty: Optional[float]
    dynamic_rating_after: Optional[float]
    rating_trend_hint: str


def _parse_date(value: str, field_name: str) -> date:
    try:
        return datetime.strptime(value, "%Y-%m-%d").date()
    except ValueError as exc:
        raise SystemExit(f"Invalid {field_name} date '{value}'. Expected YYYY-MM-DD.") from exc


def _float_or_none(value: object) -> Optional[float]:
    if value in (None, "", "-", "--", "---", "-----", "\u2014"):
        return None
    if isinstance(value, (int, float)):
        return float(value)
    text = str(value).strip()
    if not text or text.startswith("-"):
        return None
    try:
        return float(text)
    except ValueError as exc:
        raise SystemExit(f"Expected a numeric value, got '{value}'.") from exc


def _slugify(value: str) -> str:
    return re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_") or "player"


def _opponent_from_payload(payload: Dict[str, object], field_name: str) -> Opponent:
    name = str(payload.get("name", "")).strip()
    if not name:
        raise SystemExit(f"Missing {field_name}.name in input JSON.")
    return Opponent(name=name, rating=_float_or_none(payload.get("rating")))


def _player_from_payload(payload: Dict[str, object]) -> PlayerInfo:
    name = str(payload.get("name", "")).strip()
    if not name:
        raise SystemExit("Missing player.name in input JSON.")
    return PlayerInfo(
        name=name,
        location=str(payload.get("location", "")).strip(),
        ntrp_level=str(payload.get("ntrp_level", "")).strip(),
        rating_type=str(payload.get("rating_type", "-")).strip() or "-",
        dynamic_rating=_float_or_none(payload.get("dynamic_rating")),
        rating_as_of=str(payload.get("rating_as_of", "")).strip(),
        profile_url=str(payload.get("profile_url", "")).strip(),
        match_history_urls={
            int(year): str(url)
            for year, url in dict(payload.get("match_history_urls", {})).items()
        },
    )


def _match_from_payload(payload: Dict[str, object]) -> Match:
    opponents_payload = payload.get("opponents")
    if not isinstance(opponents_payload, list) or not opponents_payload:
        raise SystemExit("Each match must include a non-empty opponents array.")

    partner_payload = payload.get("partner")
    partner = None
    if isinstance(partner_payload, dict) and partner_payload:
        partner = _opponent_from_payload(partner_payload, "partner")

    result = str(payload.get("result", "")).strip().upper()
    if result not in {"W", "L"}:
        raise SystemExit(f"Invalid match result '{result}'. Expected 'W' or 'L'.")

    trend = str(payload.get("rating_trend_hint", "flat")).strip().lower() or "flat"
    if trend not in {"up", "down", "flat"}:
        trend = "flat"

    match = Match(
        date=_parse_date(str(payload.get("date", "")), "match"),
        court=str(payload.get("court", "")).strip(),
        is_singles=bool(payload.get("is_singles")),
        league=str(payload.get("league", "")).strip(),
        level=str(payload.get("level", "")).strip(),
        my_team=str(payload.get("my_team", "")).strip(),
        opponent_team=str(payload.get("opponent_team", "")).strip(),
        result=result,
        score=str(payload.get("score", "")).strip(),
        partner=partner,
        opponents=[
            _opponent_from_payload(item, "opponents[]")
            for item in opponents_payload
            if isinstance(item, dict)
        ],
        match_difficulty=_float_or_none(payload.get("match_difficulty")),
        dynamic_rating_after=_float_or_none(payload.get("dynamic_rating_after")),
        rating_trend_hint=trend,
    )
    if not match.opponents:
        raise SystemExit("Each match must include at least one parsed opponent.")
    return match


def dedupe_matches(matches: List[Match]) -> List[Match]:
    seen: Dict[Tuple[date, str, Tuple[str, ...]], Match] = {}
    for match in matches:
        key = (match.date, match.court, tuple(sorted(opponent.name for opponent in match.opponents)))
        seen[key] = match
    return sorted(seen.values(), key=lambda item: item.date, reverse=True)


def load_input_payload(input_path: Path) -> Tuple[PlayerInfo, List[Match], date, date, int]:
    try:
        payload = json.loads(input_path.read_text(encoding="utf-8"))
    except FileNotFoundError as exc:
        raise SystemExit(f"Input JSON not found: {input_path}") from exc
    except json.JSONDecodeError as exc:
        raise SystemExit(f"Invalid JSON in {input_path}: {exc}") from exc

    player_payload = payload.get("player")
    if not isinstance(player_payload, dict):
        raise SystemExit("Input JSON must contain a top-level 'player' object.")

    matches_payload = payload.get("matches", [])
    if not isinstance(matches_payload, list):
        raise SystemExit("Input JSON field 'matches' must be an array.")

    lookback_payload = payload.get("lookback", {})
    if lookback_payload and not isinstance(lookback_payload, dict):
        raise SystemExit("Input JSON field 'lookback' must be an object when provided.")

    player = _player_from_payload(player_payload)
    matches = dedupe_matches([
        _match_from_payload(item)
        for item in matches_payload
        if isinstance(item, dict)
    ])

    months = int(lookback_payload.get("months", payload.get("months", 24)))
    end_value = str(lookback_payload.get("end", payload.get("lookback_end", date.today().isoformat())))
    lookback_end = _parse_date(end_value, "lookback_end")

    start_value = lookback_payload.get("start", payload.get("lookback_start"))
    if start_value:
        lookback_start = _parse_date(str(start_value), "lookback_start")
    else:
        lookback_start = lookback_end - timedelta(days=months * 30)

    filtered_matches = [
        match for match in matches
        if lookback_start <= match.date <= lookback_end
    ]
    return player, filtered_matches, lookback_start, lookback_end, months


def _mean(values: List[float]) -> Optional[float]:
    filtered = [value for value in values if value is not None]
    return sum(filtered) / len(filtered) if filtered else None


def _median(values: List[float]) -> Optional[float]:
    filtered = sorted(value for value in values if value is not None)
    if not filtered:
        return None
    mid = len(filtered) // 2
    if len(filtered) % 2:
        return filtered[mid]
    return (filtered[mid - 1] + filtered[mid]) / 2


def _fmt(value: Optional[float], digits: int = 2) -> str:
    return f"{value:.{digits}f}" if value is not None else "\u2014"


def _slope(pairs: List[Tuple[date, float]]) -> Optional[float]:
    points = [(pair[0].toordinal(), pair[1]) for pair in pairs if pair[1] is not None]
    if len(points) < 3:
        return None
    n_points = len(points)
    mean_x = sum(point[0] for point in points) / n_points
    mean_y = sum(point[1] for point in points) / n_points
    numerator = sum((point[0] - mean_x) * (point[1] - mean_y) for point in points)
    denominator = sum((point[0] - mean_x) ** 2 for point in points)
    return numerator / denominator if denominator else None


def analyze(matches: List[Match], player_dr: Optional[float]) -> Dict[str, object]:
    total = len(matches)
    singles = [match for match in matches if match.is_singles]
    doubles = [match for match in matches if not match.is_singles]
    wins = [match for match in matches if match.result == "W"]
    losses = [match for match in matches if match.result == "L"]

    last10 = list(reversed(matches[:10]))
    form = "".join(match.result for match in last10)

    streak_count = 0
    streak_type = None
    for match in matches:
        if streak_type is None:
            streak_type = match.result
            streak_count = 1
        elif match.result == streak_type:
            streak_count += 1
        else:
            break
    streak = f"{streak_type}{streak_count}" if streak_type else "\u2014"

    difficulties = [match.match_difficulty for match in matches if match.match_difficulty is not None]
    buckets: Dict[str, int] = {}
    for difficulty in difficulties:
        bucket = f"{(int(difficulty * 4) / 4):.2f}"
        buckets[bucket] = buckets.get(bucket, 0) + 1

    def record_vs(predicate) -> Tuple[int, int]:
        wins_count = sum(
            1 for match in matches
            if match.match_difficulty is not None and predicate(match.match_difficulty) and match.result == "W"
        )
        loss_count = sum(
            1 for match in matches
            if match.match_difficulty is not None and predicate(match.match_difficulty) and match.result == "L"
        )
        return wins_count, loss_count

    if player_dr is not None:
        vs_higher = record_vs(lambda difficulty: difficulty >= player_dr + 0.25)
        vs_equal = record_vs(lambda difficulty: abs(difficulty - player_dr) < 0.25)
        vs_lower = record_vs(lambda difficulty: difficulty <= player_dr - 0.25)
    else:
        vs_higher = (0, 0)
        vs_equal = (0, 0)
        vs_lower = (0, 0)

    dr_series = [
        (match.date, match.dynamic_rating_after)
        for match in reversed(matches)
        if match.dynamic_rating_after is not None
    ]
    slope_per_day = _slope(dr_series) or 0.0
    if slope_per_day > 0.0007:
        trend = "rising"
    elif slope_per_day < -0.0007:
        trend = "falling"
    else:
        trend = "steady"

    dr_values = [value for _, value in dr_series]
    dr_first = dr_series[0][1] if dr_series else None
    dr_last = dr_series[-1][1] if dr_series else None
    dr_max = max(dr_values) if dr_values else None
    dr_min = min(dr_values) if dr_values else None

    court_counts: Dict[str, int] = {}
    for match in matches:
        court_counts[match.court] = court_counts.get(match.court, 0) + 1

    partner_stats: Dict[str, Dict[str, object]] = {}
    for match in doubles:
        if not match.partner:
            continue
        stats = partner_stats.setdefault(match.partner.name, {"matches": 0, "wins": 0, "ratings": []})
        stats["matches"] = int(stats["matches"]) + 1
        if match.result == "W":
            stats["wins"] = int(stats["wins"]) + 1
        if match.partner.rating is not None:
            ratings = stats["ratings"]
            if isinstance(ratings, list):
                ratings.append(match.partner.rating)

    partners_sorted = sorted(partner_stats.items(), key=lambda item: (-int(item[1]["matches"]), item[0]))

    top_wins = sorted(
        [match for match in wins if match.match_difficulty is not None],
        key=lambda match: match.match_difficulty,
        reverse=True,
    )[:3]
    upset_losses = sorted(
        [match for match in losses if match.match_difficulty is not None],
        key=lambda match: match.match_difficulty,
    )[:3]

    active_months = sorted({(match.date.year, match.date.month) for match in matches})
    if len(matches) >= 2:
        gaps = [(matches[index - 1].date - matches[index].date).days for index in range(1, len(matches))]
        longest_gap = max(gaps)
    else:
        longest_gap = 0

    level_counts: Dict[str, int] = {}
    for match in matches:
        key = f"{match.league.strip()} {match.level}".strip() or "Unknown"
        level_counts[key] = level_counts.get(key, 0) + 1

    jump_flags: List[Dict[str, object]] = []
    ordered_matches = list(reversed(matches))
    for index in range(1, len(ordered_matches)):
        previous = ordered_matches[index - 1]
        current = ordered_matches[index]
        if previous.dynamic_rating_after is None or current.dynamic_rating_after is None:
            continue
        delta = current.dynamic_rating_after - previous.dynamic_rating_after
        if abs(delta) > 0.3:
            jump_flags.append({
                "date": current.date,
                "delta": delta,
                "from": previous.dynamic_rating_after,
                "to": current.dynamic_rating_after,
                "match": current,
            })

    current_year = date.today().year
    level_pattern = re.compile(r"(\d\.\d)")

    def extract_ntrp(level_text: str) -> Optional[float]:
        match = level_pattern.search(level_text)
        return float(match.group(1)) if match else None

    prior_levels = set()
    current_levels = set()
    for match in matches:
        ntrp = extract_ntrp(match.level)
        if ntrp is None:
            continue
        if match.date.year < current_year:
            prior_levels.add(ntrp)
        else:
            current_levels.add(ntrp)

    stepping_down = None
    if prior_levels and current_levels:
        prior_max = max(prior_levels)
        current_max = max(current_levels)
        if current_max < prior_max:
            prior_record = [
                match.result for match in matches
                if extract_ntrp(match.level) == prior_max and match.date.year < current_year
            ]
            prior_wins = sum(1 for result in prior_record if result == "W")
            prior_losses = len(prior_record) - prior_wins
            stepping_down = (
                f"Playing {current_max} in {current_year} after competing at {prior_max} in prior season(s). "
                f"Prior record at {prior_max}: {prior_wins}W-{prior_losses}L. DR may look strong for this level."
            )

    return {
        "total": total,
        "singles_total": len(singles),
        "doubles_total": len(doubles),
        "singles_wl": (
            sum(1 for match in singles if match.result == "W"),
            sum(1 for match in singles if match.result == "L"),
        ),
        "doubles_wl": (
            sum(1 for match in doubles if match.result == "W"),
            sum(1 for match in doubles if match.result == "L"),
        ),
        "wins": len(wins),
        "losses": len(losses),
        "form": form,
        "streak": streak,
        "diff_mean": _mean(difficulties),
        "diff_median": _median(difficulties),
        "diff_min": min(difficulties) if difficulties else None,
        "diff_max": max(difficulties) if difficulties else None,
        "diff_buckets": dict(sorted(buckets.items())),
        "vs_higher": vs_higher,
        "vs_equal": vs_equal,
        "vs_lower": vs_lower,
        "dr_first": dr_first,
        "dr_last": dr_last,
        "dr_max": dr_max,
        "dr_min": dr_min,
        "dr_slope_per_day": slope_per_day,
        "trend": trend,
        "dr_series": dr_series,
        "court_counts": court_counts,
        "partners": partners_sorted,
        "top_wins": top_wins,
        "upset_losses": upset_losses,
        "months_active": len(active_months),
        "longest_gap": longest_gap,
        "level_counts": level_counts,
        "stepping_down": stepping_down,
        "jump_flags": jump_flags,
        "small_sample": total < 5,
    }


def _shade_cell(cell, color_hex: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), color_hex)
    tc_pr.append(shd)


def _set_cell_text(
    cell,
    text: str,
    *,
    bold: bool = False,
    color_hex: Optional[str] = None,
    align=WD_ALIGN_PARAGRAPH.LEFT,
    font_size: Optional[int] = None,
) -> None:
    cell.text = ""
    paragraph = cell.paragraphs[0]
    paragraph.alignment = align
    run = paragraph.add_run(text)
    run.bold = bold
    if color_hex:
        run.font.color.rgb = RGBColor.from_string(color_hex)
    if font_size is not None:
        run.font.size = Pt(font_size)


def _pct(wins: int, losses: int) -> str:
    total = wins + losses
    return f"{100 * wins / total:.1f}%" if total else "\u2014"


def _opponents_str(opponents: List[Opponent]) -> str:
    return " / ".join(
        f"{opponent.name} ({_fmt(opponent.rating)})" if opponent.rating is not None else opponent.name
        for opponent in opponents
    )


def _opponent_rating_avg(opponents: List[Opponent]) -> Optional[float]:
    ratings = [opponent.rating for opponent in opponents if opponent.rating is not None]
    return sum(ratings) / len(ratings) if ratings else None


def _add_heading(doc: Document, text: str, level: int = 1) -> None:
    paragraph = doc.add_paragraph()
    run = paragraph.add_run(text)
    run.bold = True
    sizes = {0: 22, 1: 16, 2: 13, 3: 12}
    run.font.size = Pt(sizes.get(level, 12))
    run.font.color.rgb = RGBColor.from_string(NAVY)
    if level <= 1:
        paragraph_format = paragraph.paragraph_format
        paragraph_format.space_before = Pt(12)
        paragraph_format.space_after = Pt(6)


def _add_kv_table(doc: Document, rows: List[Tuple[str, str]]) -> None:
    table = doc.add_table(rows=len(rows), cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.autofit = True
    for index, (key, value) in enumerate(rows):
        _set_cell_text(table.rows[index].cells[0], key, bold=True, font_size=10)
        _set_cell_text(table.rows[index].cells[1], value, font_size=10)
        _shade_cell(table.rows[index].cells[0], LIGHT_BLUE)


def _add_header_row(table, headers: List[str]) -> None:
    for cell, header in zip(table.rows[0].cells, headers):
        _set_cell_text(
            cell,
            header,
            bold=True,
            color_hex=WHITE,
            font_size=10,
            align=WD_ALIGN_PARAGRAPH.CENTER,
        )
        _shade_cell(cell, NAVY)


def render_docx(
    player: PlayerInfo,
    matches: List[Match],
    analysis: Dict[str, object],
    lookback_start: date,
    lookback_end: date,
    months: int,
    out_path: Path,
) -> None:
    doc = Document()

    for section in doc.sections:
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)
        section.top_margin = Inches(0.7)
        section.bottom_margin = Inches(0.7)

    style = doc.styles["Normal"]
    style.font.name = "Calibri"
    style.font.size = Pt(10)

    title = doc.add_paragraph()
    title_run = title.add_run(f"Player Scout: {player.name}")
    title_run.bold = True
    title_run.font.size = Pt(20)
    title_run.font.color.rgb = RGBColor.from_string(NAVY)

    header_bits = []
    if player.location:
        header_bits.append(f"Location: {player.location}")
    if player.ntrp_level:
        header_bits.append(f"NTRP: {player.ntrp_level} {player.rating_type}".strip())
    if player.dynamic_rating is not None:
        dr_text = f"Dynamic Rating: {_fmt(player.dynamic_rating, 4)}"
        if player.rating_as_of:
            dr_text += f" (as of {player.rating_as_of})"
        header_bits.append(dr_text)
    if header_bits:
        subtitle = doc.add_paragraph(" · ".join(header_bits))
        subtitle.runs[0].font.size = Pt(11)

    meta = doc.add_paragraph()
    meta.add_run(
        f"Lookback: {lookback_start.isoformat()} -> {lookback_end.isoformat()} ({months} months)"
    ).font.size = Pt(9)
    if player.profile_url:
        meta.add_run("\n").font.size = Pt(9)
        source_run = meta.add_run(f"Source: {player.profile_url}")
        source_run.font.size = Pt(9)
        source_run.font.color.rgb = RGBColor.from_string(MID_BLUE)

    if player.rating_type == "S":
        warning = doc.add_paragraph()
        warning_run = warning.add_run(
            "Self-rated: may play above listed NTRP. Treat the listed level as a floor, not a ceiling."
        )
        warning_run.bold = True
        warning_run.font.color.rgb = RGBColor.from_string("C00000")

    _add_heading(doc, "TL;DR", level=1)
    bullets = [
        f"{analysis['total']} matches in window ({analysis['singles_total']} singles, {analysis['doubles_total']} doubles), {analysis['wins']}-{analysis['losses']} overall.",
    ]
    if player.rating_type == "S":
        bullets.append("Self-rated player: recent results may understate the level they can bring on the day.")
    if analysis["diff_mean"] is not None:
        bullets.append(
            f"Typical opponent strength {_fmt(analysis['diff_mean'])} with range {_fmt(analysis['diff_min'])}-{_fmt(analysis['diff_max'])}."
        )
    if analysis["dr_first"] is not None and analysis["dr_last"] is not None:
        delta = float(analysis["dr_last"]) - float(analysis["dr_first"])
        sign = "+" if delta >= 0 else ""
        bullets.append(
            f"Dynamic rating {analysis['trend']}: {_fmt(analysis['dr_first'], 3)} -> {_fmt(analysis['dr_last'], 3)} ({sign}{delta:.3f})."
        )
    if analysis["form"]:
        bullets.append(f"Last {len(str(analysis['form']))} matches: {analysis['form']} (streak {analysis['streak']}).")
    if analysis["small_sample"]:
        bullets.append("Small sample: fewer than 5 matches in the window, so trend and split conclusions are low-confidence.")
    jump_flags = analysis["jump_flags"]
    if jump_flags:
        first_jump = jump_flags[0]
        delta = float(first_jump["delta"])
        bullets.append(
            f"Rating jump flag: {first_jump['date'].isoformat()} moved from {_fmt(first_jump['from'], 3)} to {_fmt(first_jump['to'], 3)} ({delta:+.3f}), which may indicate a self-rate bump or data anomaly."
        )
    for bullet in bullets[:5]:
        doc.add_paragraph(str(bullet), style="List Bullet")

    if analysis["stepping_down"]:
        stepping_warning = doc.add_paragraph()
        stepping_run = stepping_warning.add_run(f"Stepping down: {analysis['stepping_down']}")
        stepping_run.bold = True
        stepping_run.font.color.rgb = RGBColor.from_string("C00000")

    _add_heading(doc, "Volume & Activity", level=1)
    _add_kv_table(
        doc,
        [
            ("Total matches", str(analysis["total"])),
            ("Singles / Doubles", f"{analysis['singles_total']} / {analysis['doubles_total']}"),
            ("Active months", str(analysis["months_active"])),
            ("Longest gap between matches", f"{analysis['longest_gap']} days"),
        ],
    )
    if analysis["level_counts"]:
        level_paragraph = doc.add_paragraph()
        level_paragraph.add_run("Play mix by league/level:").bold = True
        for key, value in sorted(analysis["level_counts"].items(), key=lambda item: -item[1]):
            doc.add_paragraph(f"{key}: {value}", style="List Bullet")

    _add_heading(doc, "Results", level=1)
    results_table = doc.add_table(rows=4, cols=4)
    results_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _add_header_row(results_table, ["Split", "W", "L", "Win %"])
    singles_w, singles_l = analysis["singles_wl"]
    doubles_w, doubles_l = analysis["doubles_wl"]
    for index, (label, wins_count, losses_count) in enumerate(
        [
            ("Overall", analysis["wins"], analysis["losses"]),
            ("Singles", singles_w, singles_l),
            ("Doubles", doubles_w, doubles_l),
        ],
        start=1,
    ):
        cells = results_table.rows[index].cells
        _set_cell_text(cells[0], str(label), font_size=10)
        _set_cell_text(cells[1], str(wins_count), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(cells[2], str(losses_count), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(cells[3], _pct(int(wins_count), int(losses_count)), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    form_paragraph = doc.add_paragraph()
    form_paragraph.add_run(f"Last {len(str(analysis['form']))} matches: ").bold = True
    form_paragraph.add_run(str(analysis["form"]) or "\u2014")
    streak_paragraph = doc.add_paragraph()
    streak_paragraph.add_run("Current streak: ").bold = True
    streak_paragraph.add_run(str(analysis["streak"]))

    _add_heading(doc, "Opponent Strength", level=1)
    _add_kv_table(
        doc,
        [
            ("Mean match difficulty", _fmt(analysis["diff_mean"])),
            ("Median", _fmt(analysis["diff_median"])),
            ("Range", f"{_fmt(analysis['diff_min'])} -> {_fmt(analysis['diff_max'])}"),
        ],
    )
    if player.dynamic_rating is not None:
        relative_paragraph = doc.add_paragraph()
        relative_paragraph.add_run(f"Relative to current DR ({_fmt(player.dynamic_rating, 3)}):").bold = True
        relative_table = doc.add_table(rows=4, cols=4)
        _add_header_row(relative_table, ["Bucket", "W", "L", "Win %"])
        for index, (label, record) in enumerate(
            [
                ("Tougher (opp >= DR+0.25)", analysis["vs_higher"]),
                ("Peer (within +/-0.25)", analysis["vs_equal"]),
                ("Easier (opp <= DR-0.25)", analysis["vs_lower"]),
            ],
            start=1,
        ):
            wins_count, losses_count = record
            cells = relative_table.rows[index].cells
            _set_cell_text(cells[0], str(label), font_size=10)
            _set_cell_text(cells[1], str(wins_count), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(cells[2], str(losses_count), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(cells[3], _pct(int(wins_count), int(losses_count)), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    if analysis["diff_buckets"]:
        bucket_paragraph = doc.add_paragraph()
        bucket_paragraph.add_run("Opponent difficulty distribution (0.25 buckets):").bold = True
        bucket_table = doc.add_table(rows=1 + len(analysis["diff_buckets"]), cols=2)
        _add_header_row(bucket_table, ["Bucket", "Count"])
        for index, (bucket, count) in enumerate(analysis["diff_buckets"].items(), start=1):
            cells = bucket_table.rows[index].cells
            _set_cell_text(cells[0], str(bucket), font_size=10)
            _set_cell_text(cells[1], str(count), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    _add_heading(doc, "Rating Trend", level=1)
    if analysis["dr_first"] is not None and analysis["dr_last"] is not None:
        n_points = len(analysis["dr_series"])
        delta = float(analysis["dr_last"]) - float(analysis["dr_first"])
        sign = "+" if delta >= 0 else ""
        summary = doc.add_paragraph()
        summary_run = summary.add_run(
            f"DR Trend: {_fmt(analysis['dr_first'], 3)} -> {_fmt(analysis['dr_last'], 3)} ({sign}{delta:.3f} across {n_points} post-match samples, {analysis['trend']})"
        )
        summary_run.bold = True
    _add_kv_table(
        doc,
        [
            ("First DR in window", _fmt(analysis["dr_first"], 4)),
            ("Last DR in window", _fmt(analysis["dr_last"], 4)),
            ("Max / Min", f"{_fmt(analysis['dr_max'], 4)} / {_fmt(analysis['dr_min'], 4)}"),
            ("Trend", f"{analysis['trend']} (slope {float(analysis['dr_slope_per_day']):+.5f}/day)"),
        ],
    )
    if analysis["small_sample"]:
        caveat = doc.add_paragraph()
        caveat_run = caveat.add_run("Small sample caveat: fewer than 5 rating points in-window, so treat the slope as directional only.")
        caveat_run.italic = True
    if jump_flags:
        jumps_paragraph = doc.add_paragraph()
        jumps_paragraph.add_run("Rating jump checkpoints:").bold = True
        for jump in jump_flags:
            doc.add_paragraph(
                f"{jump['date'].isoformat()}: {_fmt(jump['from'], 3)} -> {_fmt(jump['to'], 3)} ({float(jump['delta']):+.3f})",
                style="List Bullet",
            )

    _add_heading(doc, "Court & Partners", level=1)
    if analysis["court_counts"]:
        court_table = doc.add_table(rows=1 + len(analysis["court_counts"]), cols=2)
        _add_header_row(court_table, ["Court", "Matches"])
        for index, (court, count) in enumerate(sorted(analysis["court_counts"].items()), start=1):
            cells = court_table.rows[index].cells
            _set_cell_text(cells[0], str(court), font_size=10)
            _set_cell_text(cells[1], str(count), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    if analysis["partners"]:
        partner_intro = doc.add_paragraph()
        partner_intro.add_run("Most frequent doubles partners:").bold = True
        top_partners = analysis["partners"][:5]
        partner_table = doc.add_table(rows=1 + len(top_partners), cols=5)
        _add_header_row(partner_table, ["Partner", "Matches", "Wins", "Win %", "Avg Rating"])
        for index, (name, stats) in enumerate(top_partners, start=1):
            ratings = stats["ratings"] if isinstance(stats["ratings"], list) else []
            avg_rating = sum(ratings) / len(ratings) if ratings else None
            matches_count = int(stats["matches"])
            wins_count = int(stats["wins"])
            cells = partner_table.rows[index].cells
            _set_cell_text(cells[0], str(name), font_size=10)
            _set_cell_text(cells[1], str(matches_count), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(cells[2], str(wins_count), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(cells[3], _pct(wins_count, matches_count - wins_count), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)
            _set_cell_text(cells[4], _fmt(avg_rating), font_size=10, align=WD_ALIGN_PARAGRAPH.RIGHT)

    if analysis["top_wins"] or analysis["upset_losses"]:
        _add_heading(doc, "Notable Matches", level=1)
        if analysis["top_wins"]:
            wins_intro = doc.add_paragraph()
            wins_intro.add_run("Biggest wins (by opponent strength):").bold = True
            for match in analysis["top_wins"]:
                doc.add_paragraph(
                    f"{match.date.isoformat()} · {match.court} · vs {_opponents_str(match.opponents)} · {match.score} · diff {_fmt(match.match_difficulty)}",
                    style="List Bullet",
                )
        if analysis["upset_losses"]:
            losses_intro = doc.add_paragraph()
            losses_intro.add_run("Losses to notably weaker opponents:").bold = True
            for match in analysis["upset_losses"]:
                doc.add_paragraph(
                    f"{match.date.isoformat()} · {match.court} · vs {_opponents_str(match.opponents)} · {match.score} · diff {_fmt(match.match_difficulty)}",
                    style="List Bullet",
                )

    _add_heading(doc, "Full Match Log", level=1)
    headers = ["Date", "Type", "Court", "Opp Team", "Opponent(s)", "Opp Rating", "Score", "W/L", "Post DR", "DR Δ"]
    log_table = doc.add_table(rows=1 + len(matches), cols=len(headers))
    _add_header_row(log_table, headers)
    ordered_matches = list(reversed(matches))
    previous_dr_map: Dict[int, Optional[float]] = {}
    for index, match in enumerate(ordered_matches):
        previous_dr_map[index] = ordered_matches[index - 1].dynamic_rating_after if index > 0 else None

    for index, match in enumerate(matches, start=1):
        ordered_index = len(matches) - index
        previous_dr = previous_dr_map.get(ordered_index)
        current_dr = match.dynamic_rating_after
        delta_text = f"{current_dr - previous_dr:+.2f}" if current_dr is not None and previous_dr is not None else "\u2014"

        avg_rating = _opponent_rating_avg(match.opponents)
        cells = log_table.rows[index].cells
        _set_cell_text(cells[0], match.date.isoformat(), font_size=9)
        _set_cell_text(cells[1], "S" if match.is_singles else "D", font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[2], match.court, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER)
        _set_cell_text(cells[3], match.opponent_team, font_size=9)
        _set_cell_text(cells[4], _opponents_str(match.opponents), font_size=9)
        _set_cell_text(cells[5], _fmt(avg_rating), font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(cells[6], match.score, font_size=9)
        _set_cell_text(cells[7], match.result, font_size=9, align=WD_ALIGN_PARAGRAPH.CENTER, bold=True)
        _set_cell_text(cells[8], _fmt(current_dr, 4), font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        _set_cell_text(cells[9], delta_text, font_size=9, align=WD_ALIGN_PARAGRAPH.RIGHT)
        if match.result == "W":
            _shade_cell(cells[7], WIN_GREEN)
        elif match.result == "L":
            _shade_cell(cells[7], LOSS_RED)

    if not matches:
        no_data = doc.add_paragraph()
        no_data_run = no_data.add_run("No matches in the selected lookback window. The report reflects profile data only.")
        no_data_run.italic = True

    footer = doc.add_paragraph()
    footer_run = footer.add_run(
        f"Scouted {datetime.now().strftime('%Y-%m-%d %H:%M')} from pre-fetched public tennisrecord.com data."
    )
    footer_run.italic = True
    footer_run.font.size = Pt(9)
    footer_run.font.color.rgb = RGBColor.from_string("7A7A7A")

    doc.save(str(out_path))


def print_inline_summary(player: PlayerInfo, analysis: Dict[str, object]) -> None:
    dr_text = _fmt(player.dynamic_rating, 2) if player.dynamic_rating is not None else "\u2014"
    rating_type = player.rating_type if player.rating_type != "-" else "?"
    print(f"  Player:  {player.name}" + (f" ({player.location})" if player.location else ""))
    print(
        f"  Rating:  {player.ntrp_level or '—'} {rating_type} | DR: {dr_text}"
        + (f" as of {player.rating_as_of}" if player.rating_as_of else "")
    )
    if analysis["total"] == 0:
        print("  Matches: No match history in the selected lookback window.")
        return
    print(
        f"  Record:  {analysis['wins']}W-{analysis['losses']}L "
        f"({analysis['total']} matches, {analysis['singles_total']}S/{analysis['doubles_total']}D)"
    )
    print(f"  Form:    {analysis['form'] or '—'} (streak: {analysis['streak']})")
    if analysis["dr_first"] is not None and analysis["dr_last"] is not None:
        delta = float(analysis["dr_last"]) - float(analysis["dr_first"])
        sign = "+" if delta >= 0 else ""
        print(
            f"  DR Trend: {analysis['trend']} ({_fmt(analysis['dr_first'], 3)} -> {_fmt(analysis['dr_last'], 3)}, {sign}{delta:.3f})"
        )
    courts = ", ".join(f"{court}x{count}" for court, count in sorted(analysis["court_counts"].items()))
    if courts:
        print(f"  Courts:  {courts}")
    if analysis["jump_flags"]:
        first_jump = analysis["jump_flags"][0]
        print(
            f"  Jump:    {first_jump['date'].isoformat()} {_fmt(first_jump['from'], 3)} -> {_fmt(first_jump['to'], 3)} ({float(first_jump['delta']):+.3f})"
        )


def main(argv: Optional[List[str]] = None) -> int:
    parser = argparse.ArgumentParser(description=__doc__)
    parser.add_argument("--input", required=True, help="Path to normalized player JSON data")
    parser.add_argument("--output", default=None, help="Output filename (placed under reports/)")
    parser.add_argument("--workspace", default=None, help="Workspace root (defaults to repo root)")
    parser.add_argument("--summary", action="store_true", help="Print a quick inline summary instead of generating a .docx")
    args = parser.parse_args(argv)

    workspace = Path(args.workspace).resolve() if args.workspace else Path(__file__).resolve().parents[3]
    reports_dir = workspace / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    input_path = Path(args.input).resolve()
    player, matches, lookback_start, lookback_end, months = load_input_payload(input_path)
    analysis = analyze(matches, player.dynamic_rating)

    if args.summary:
        print_inline_summary(player, analysis)
        return 0

    if args.output:
        filename = Path(args.output).name
        if not filename.lower().endswith(".docx"):
            filename += ".docx"
    else:
        loc_slug = _slugify(player.location.split(",")[0]) if player.location else "loc"
        name_bits = player.name.split()
        first_slug = _slugify(name_bits[0]) if name_bits else "player"
        last_slug = _slugify(" ".join(name_bits[1:])) if len(name_bits) > 1 else "unknown"
        filename = f"player_{first_slug}_{last_slug}_{loc_slug}_{lookback_end.strftime('%Y%m%d')}.docx"
    output_path = reports_dir / filename

    render_docx(player, matches, analysis, lookback_start, lookback_end, months, output_path)

    print(f"[usta-player-scout] Wrote {output_path}")
    print(
        f"[usta-player-scout] {analysis['total']} matches in {months}m window "
        f"({analysis['singles_total']}S / {analysis['doubles_total']}D); "
        f"record {analysis['wins']}-{analysis['losses']}"
    )
    return 0


if __name__ == "__main__":
    sys.exit(main())
