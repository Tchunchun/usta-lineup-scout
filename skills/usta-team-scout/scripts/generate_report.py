import argparse
import difflib
import json
import re
import subprocess
import unicodedata
from concurrent.futures import ThreadPoolExecutor, as_completed
from collections import Counter
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from urllib.parse import quote, urljoin

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_URL = "https://www.tennisrecord.com"

NAVY = "1F3864"
LIGHT_BLUE = "D6E4F0"
MID_BLUE = "2E75B6"
WIN_GREEN = "E2EFDA"
LOSS_RED = "FCE4D6"
WARNING_YELLOW = "FFF2CC"
WHITE = "FFFFFF"


@dataclass
class Player:
    name: str
    href: str
    location: str
    ntrp: str
    season_record: str
    local_singles: str
    local_doubles: str
    local_record: str
    dr: Optional[float]
    rating_type: str = "—"


@dataclass
class MatchCourt:
    court: str
    team_players: List[Tuple[str, Optional[float], str, str]]
    opponent_players: List[Tuple[str, Optional[float], str, str]]
    score: str
    result: str


@dataclass
class MatchReport:
    date: str
    site: str
    team_name: str
    opponent: str
    final_score: str
    team_won_match: bool
    courts: List[MatchCourt]


def fetch(url: str) -> str:
    safe_url = quote(url, safe=":/?&=%")
    return subprocess.check_output(
        ["curl", "-L", "--silent", "--connect-timeout", "30", "--max-time", "60", safe_url],
        text=True,
    )


def soup_from_url(url: str) -> BeautifulSoup:
    return BeautifulSoup(fetch(url), "lxml")


def normalize_space(text: str) -> str:
    return " ".join(text.split())


def canonicalize_name(name: str) -> str:
    normalized = unicodedata.normalize("NFKD", normalize_space(name))
    ascii_name = "".join(ch for ch in normalized if not unicodedata.combining(ch))
    return re.sub(r"[^a-z0-9]+", "", ascii_name.lower())


def resolve_roster_player(name: str, roster_by_key: Dict[str, Player]) -> Optional[Player]:
    key = canonicalize_name(name)
    direct = roster_by_key.get(key)
    if direct:
        return direct

    matches = difflib.get_close_matches(key, roster_by_key.keys(), n=1, cutoff=0.92)
    if matches:
        return roster_by_key[matches[0]]

    return None


def slugify_filename(value: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_")
    return slug or "Scouting_Report"


def report_date() -> str:
    return date.today().strftime("%Y%m%d")


def infer_team_level(league_label: str, roster: List[Player]) -> str:
    league_match = re.search(r"\b(\d\.\d)\b", league_label)
    if league_match:
        return league_match.group(1)
    for player in roster:
        if player.ntrp:
            return player.ntrp
    return "unknown"


def fuzzy_team_name_candidates(team_name: str) -> List[str]:
    """Return alternate spellings to try when the exact name fails."""
    candidates = [team_name]
    # space ↔ hyphen swap
    if " " in team_name:
        candidates.append(team_name.replace(" ", "-"))
    if "-" in team_name:
        candidates.append(team_name.replace("-", " "))
    # collapse repeated separators
    collapsed = re.sub(r"[-\s]{2,}", "-", team_name)
    if collapsed not in candidates:
        candidates.append(collapsed)
    # strip trailing NTRP token (e.g. "Team 3.0" → "Team")
    stripped = re.sub(r"\s+\d\.\d$", "", team_name).strip()
    if stripped and stripped not in candidates:
        candidates.append(stripped)
    return candidates


def build_team_url(team_name: str, year: int, suffix: Optional[int] = None) -> str:
    url = f"{BASE_URL}/adult/teamprofile.aspx?year={year}&teamname={quote(team_name, safe='')}"
    if suffix is not None:
        url += f"&s={suffix}"
    return url


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: Optional[str] = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(9)


def fetch_wildcard_prior_season(player_href: str, current_year: int) -> str:
    """
    For a wildcard player (zero current-season matches), fetch their prior season
    summary from TennisRecord and return a compact string like '2025: 8W-3L, D2/D3'.
    Returns empty string if no prior history found.
    """
    prior_year = current_year - 1
    # Build match history URL for prior year
    # href looks like /adult/profile.aspx?playername=First+Last&s=N
    # Match history URL: /adult/matchhistory.aspx?year=YYYY&playername=First+Last&s=N
    s_match = re.search(r"&s=(\d+)", player_href)
    name_match = re.search(r"playername=([^&]+)", player_href)
    if not name_match:
        return ""
    playername_param = name_match.group(1)
    history_url = f"{BASE_URL}/adult/matchhistory.aspx?year={prior_year}&playername={playername_param}"
    if s_match:
        history_url += f"&s={s_match.group(1)}"

    try:
        html = fetch(history_url)
    except Exception:
        return ""

    soup = BeautifulSoup(html, "lxml")
    rows = soup.find_all("tr")
    wins, losses = 0, 0
    courts_seen: set = set()
    for row in rows:
        cells = row.find_all("td")
        if len(cells) < 4:
            continue
        result_text = normalize_space(cells[-1].get_text(" ", strip=True))
        court_text = normalize_space(cells[0].get_text(" ", strip=True))
        if result_text in ("W", "Win"):
            wins += 1
        elif result_text in ("L", "Loss"):
            losses += 1
        for label in ["S1", "S2", "D1", "D2", "D3", "Singles #1", "Singles #2", "Doubles #1", "Doubles #2", "Doubles #3"]:
            if label in court_text:
                short = label if len(label) == 2 else ("S1" if "Singles #1" in label else "S2" if "Singles #2" in label else "D1" if "Doubles #1" in label else "D2" if "Doubles #2" in label else "D3")
                courts_seen.add(short)

    if wins + losses == 0:
        return ""

    courts_str = "/".join(sorted(courts_seen)) if courts_seen else "?"
    return f"{prior_year}: {wins}W-{losses}L, {courts_str}"


def fetch_player_info_by_name(name: str, year: int) -> Tuple[Optional[float], str, str]:
    """
    Look up a player's Dynamic Rating and USTA rating from their TennisRecord profile by name.
    Returns (dr, rating_type, usta_rating) e.g. (2.9155, 'C', '3.0C').
    """
    url = f"{BASE_URL}/adult/profile.aspx?playername={quote(name, safe='')}&year={year}"
    try:
        html = fetch(url)
        dr: Optional[float] = None
        dr_m = re.search(r'<span style="font-weight:bold;">(\d\.\d{3,})\s*</span>', html)
        if dr_m:
            dr = float(dr_m.group(1))
        ntrp_m = re.search(r'<span style="font-weight:bold;">(\d\.\d)\s+([A-Z])</span>', html)
        if ntrp_m:
            ntrp, suffix = ntrp_m.groups()
            rating_type = suffix if suffix in {"C", "S"} else "—"
            return dr, rating_type, f"{ntrp}{rating_type}"
        return dr, "—", "—"
    except Exception:
        return None, "—", "—"


def parse_player_rating(profile_href: str) -> Tuple[str, str]:
    url = urljoin(BASE_URL, profile_href)
    html = fetch(url)
    match = re.search(r"<span style=\"font-weight:bold;\">(\d\.\d)\s+([A-Z])</span>", html)
    if match:
        ntrp, rating_suffix = match.groups()
        rating_type = rating_suffix if rating_suffix in {"C", "S"} else "—"
        return rating_type, f"{ntrp}{rating_type}"
    return "—", "—"


def parse_dr(value: str) -> Optional[float]:
    try:
        return float(value)
    except ValueError:
        return None


def parse_players_from_cell(cell) -> List[Tuple[str, Optional[float]]]:
    players: List[Tuple[str, Optional[float]]] = []
    for anchor in cell.find_all("a", href=lambda h: h and "/adult/profile.aspx?playername=" in h):
        name = normalize_space(anchor.get_text(" ", strip=True))
        trailing = anchor.next_sibling or ""
        rating_match = re.search(r"\(([-0-9.]+|-----)\)", str(trailing))
        dr = None if not rating_match or rating_match.group(1) == "-----" else float(rating_match.group(1))
        players.append((name, dr))
    return players


def format_player_line(players: List[Tuple[str, Optional[float], str, str]]) -> str:
    formatted = []
    for name, _, _, usta_rating in players:
        formatted.append(f"{name} ({usta_rating})")
    return " / ".join(formatted)


def parse_team_page(team_name: str, year: int, forced_suffix: Optional[int] = None) -> Tuple[str, str, str, List[Player], List[Tuple[str, str, str, str]]]:
    soup = None
    resolved_team_name_found = team_name
    suffixes = [forced_suffix] if forced_suffix is not None else [None, 1, 2, 3, 4, 5]
    name_candidates = [team_name] if forced_suffix is not None else fuzzy_team_name_candidates(team_name)

    for candidate_name in name_candidates:
        for suffix in suffixes:
            page = soup_from_url(build_team_url(candidate_name, year, suffix))
            tables = page.find_all("table")
            if len(tables) < 5:
                continue
            meta_rows = tables[1].find_all("tr")
            if len(meta_rows) < 3:
                continue
            # When a forced suffix is given, accept any valid page regardless of resolved name
            if forced_suffix is not None:
                soup = page
                resolved_team_name_found = normalize_space(meta_rows[2].get_text(" ", strip=True))
                break
            resolved_name = normalize_space(meta_rows[2].get_text(" ", strip=True))
            if resolved_name == candidate_name:
                soup = page
                resolved_team_name_found = resolved_name
                break
        if soup is not None:
            break

    if soup is None:
        raise ValueError(f"Unable to locate team page for '{team_name}' in {year}. Tried: {name_candidates}")

    tables = soup.find_all("table")
    meta_rows = tables[1].find_all("tr")
    league_label = normalize_space(meta_rows[0].get_text(" ", strip=True))
    flight_label = normalize_space(meta_rows[1].get_text(" ", strip=True))
    resolved_team_name = normalize_space(meta_rows[2].get_text(" ", strip=True))

    roster: List[Player] = []
    roster_table = tables[2]
    for tr in roster_table.find_all("tr")[1:]:
        cells = tr.find_all("td")
        if len(cells) != 9:
            continue
        link = tr.find("a", href=lambda h: h and "/adult/profile.aspx?playername=" in h)
        if not link:
            continue
        texts = [normalize_space(td.get_text(" ", strip=True)) for td in cells]
        roster.append(
            Player(
                name=normalize_space(link.get_text(" ", strip=True)),
                href=link.get("href"),
                location=texts[1],
                ntrp=texts[2],
                season_record=texts[3],
                local_singles=texts[4],
                local_doubles=texts[5],
                local_record=texts[6],
                dr=parse_dr(texts[7]),
            )
        )

    schedule_table = tables[4]
    matches: List[Tuple[str, str, str, str]] = []
    for tr in schedule_table.find_all("tr")[1:]:
        cells = tr.find_all("td")
        if len(cells) != 5:
            continue
        match_date = normalize_space(cells[0].get_text(" ", strip=True))
        opponent = normalize_space(cells[2].get_text(" ", strip=True))
        result = normalize_space(cells[4].get_text(" ", strip=True))
        link = cells[4].find("a", href=lambda h: h and "/adult/matchresults.aspx?year=" in h)
        if link:
            matches.append((match_date, opponent, result, link.get("href")))
    return league_label, flight_label, resolved_team_name, roster, matches


def parse_match(
    match_href: str,
    rating_types: Dict[str, str],
    usta_ratings: Dict[str, str],
    team_name: str,
) -> MatchReport:
    soup = soup_from_url(urljoin(BASE_URL, match_href))
    tables = soup.find_all("table")

    details_text = normalize_space(tables[1].get_text(" ", strip=True))
    date_match = re.search(r"Scheduled Date:\s*([0-9/]+)", details_text)
    site_match = re.search(r"Match Site:\s*(.*?)\s*Match Criteria:", details_text)

    summary_rows = tables[2].find_all("tr")[1:]
    summary = []
    for row in summary_rows:
        cells = row.find_all("td")
        if len(cells) >= 4:
            summary.append([normalize_space(cell.get_text(" ", strip=True)) for cell in cells])
    team_is_left = summary[0][0] == team_name
    opponent = summary[1][0] if team_is_left else summary[0][0]
    wins = summary[0][1] if team_is_left else summary[1][1]
    losses = summary[1][1] if team_is_left else summary[0][1]
    team_won_match = int(wins) > int(losses)

    court_labels = ["Singles #1", "Singles #2", "Doubles #1", "Doubles #2", "Doubles #3"]
    court_keys = ["S1", "S2", "D1", "D2", "D3"]
    courts: List[MatchCourt] = []

    for label, key in zip(court_labels, court_keys):
        marker = soup.find(string=lambda s: isinstance(s, str) and label in s)
        if not marker:
            continue
        court_table = marker.parent.find_next("table")
        rows = court_table.find_all("tr")
        if len(rows) < 2:
            continue
        row = rows[1]
        cells = row.find_all("td")
        left_players = parse_players_from_cell(cells[0])
        right_players = parse_players_from_cell(cells[-1])
        score = normalize_space(cells[3].get_text(" ", strip=True)).replace(" - ", "-")
        row_html = str(row)
        winner_side = "left" if "arrowhead_right.png" in row_html else "right"
        team_side = "left" if team_is_left else "right"
        result = "W" if winner_side == team_side else "L"

        team_players_raw = left_players if team_side == "left" else right_players
        opp_players_raw = right_players if team_side == "left" else left_players
        team_players = [
            (name, dr, rating_types.get(name, "—"), usta_ratings.get(name, "—"))
            for name, dr in team_players_raw
        ]
        opponent_players = [
            (name, dr, rating_types.get(name, "—"), usta_ratings.get(name, "—"))
            for name, dr in opp_players_raw
        ]

        courts.append(
            MatchCourt(
                court=key,
                team_players=team_players,
                opponent_players=opponent_players,
                score=score,
                result=result,
            )
        )

    return MatchReport(
        date=date_match.group(1) if date_match else "",
        site=site_match.group(1) if site_match else "",
        team_name=team_name,
        opponent=opponent,
        final_score=f"{wins}-{losses}",
        team_won_match=team_won_match,
        courts=courts,
    )


@dataclass
class CourtStrategy:
    likely_lineup: str
    analysis: str


def _format_player(name: str, dr: Optional[float], rating_type: str) -> str:
    dr_str = f"{dr:.2f}" if dr is not None else "-----"
    suffix = " ⚠" if rating_type == "S" else ""
    return f"{name} ({dr_str}, {rating_type}){suffix}"


def build_strategy(
    matches: List[MatchReport],
    wildcards: List[str],
    roster: List[Player],
    team_name: str,
) -> Dict[str, CourtStrategy]:
    strategy: Dict[str, CourtStrategy] = {}
    court_map: Dict[str, List[Tuple[int, MatchCourt, MatchReport]]] = {
        c: [] for c in ("S1", "S2", "D1", "D2", "D3")
    }
    for idx, match in enumerate(matches, 1):
        for court in match.courts:
            court_map[court.court].append((idx, court, match))

    for court, entries in court_map.items():
        if not entries:
            strategy[court] = CourtStrategy(
                likely_lineup="No data yet",
                analysis="No completed sample for this court.",
            )
            continue

        # --- Lineup prediction ---

        player_freq: Counter = Counter()
        for _, entry, _ in entries:
            for name, dr, rating_type, _ in entry.team_players:
                player_freq[name] += 1

        is_doubles = court.startswith("D")
        top_players = player_freq.most_common(4 if is_doubles else 2)

        if is_doubles:
            # find most common *pair* by grouping team_players per entry
            pair_freq: Counter = Counter()
            for _, entry, _ in entries:
                pair_key = " / ".join(
                    sorted(n for n, _, _, _ in entry.team_players)
                )
                pair_freq[pair_key] += 1
            best_pair = pair_freq.most_common(1)[0][0]
            lineup_names = best_pair.split(" / ")
        else:
            lineup_names = [tp[0] for tp in top_players[:1]]

        lineup_parts = []
        for ln in lineup_names:
            matched = next(
                (
                    (name, dr, rt)
                    for _, entry, _ in entries
                    for name, dr, rt, _ in entry.team_players
                    if name == ln
                ),
                None,
            )
            if matched:
                lineup_parts.append(_format_player(*matched))
            else:
                lineup_parts.append(ln)
        likely_lineup = "\n".join(lineup_parts) if is_doubles else (lineup_parts[0] if lineup_parts else "TBD")

        # If there are alternates, note them
        alternate_names = {n for n, _ in top_players} - set(lineup_names)
        if alternate_names:
            alt_strs = []
            for an in sorted(alternate_names):
                matched = next(
                    (
                        (name, dr, rt)
                        for _, entry, _ in entries
                        for name, dr, rt, _ in entry.team_players
                        if name == an
                    ),
                    None,
                )
                alt_strs.append(_format_player(*matched) if matched else an)
            likely_lineup += "\nor " + " / ".join(alt_strs)

        # --- Analysis per match ---
        analysis_parts = []
        for idx, entry, match in entries:
            team_str = " / ".join(
                _format_player(n, d, rt)
                for n, d, rt, _ in entry.team_players
            )
            analysis_parts.append(
                f"Match {idx} vs {match.opponent}: {team_str} — "
                f"{entry.result} {entry.score}."
            )

        # Add court-level insight
        wins = sum(1 for _, e, _ in entries if e.result == "W")
        losses = len(entries) - wins
        if wins == len(entries):
            analysis_parts.append(f"Undefeated on this court ({wins}-0). Expect their strongest players here.")
        elif losses == len(entries):
            analysis_parts.append(f"Winless on this court (0-{losses}). This is a vulnerable line.")
        else:
            analysis_parts.append(f"Record: {wins}-{losses} on this court.")

        # Flag self-rated
        s_rated = sorted(set(
            n for _, e, _ in entries
            for n, _, rt, _ in e.team_players
            if rt == "S"
        ))
        if s_rated:
            analysis_parts.append(
                f"Self-rated: {', '.join(s_rated)} ⚠ — true level unknown."
            )

        strategy[court] = CourtStrategy(
            likely_lineup=likely_lineup,
            analysis=" ".join(analysis_parts),
        )
    return strategy


def add_title(document: Document, team_name: str, league_label: str, flight_label: str, season_record: str, most_recent_match: Optional[str] = None) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{team_name} Scouting Report")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    lines = [
        f"League: {league_label}",
        f"Flight: {flight_label}",
        f"Prepared: {date.today().isoformat()}",
        f"Most Recent Match: {most_recent_match}" if most_recent_match else None,
        f"Completed Match Record: {season_record}",
    ]
    for line in lines:
        if line is None:
            continue
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(10)


def add_legend(document: Document, team_name: str = "") -> None:
    document.add_heading("Legend", level=2)
    label = team_name if team_name else "home team"
    for line in [
        "◆ = DR 3.0+ player",
        "⚠ = Self-rated player (S)",
        "★ = Player on roster with no completed-match appearance yet",
        f"Green result cells = {label} win",
        f"Red result cells = {label} loss",
    ]:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(line)


def add_roster_table(document: Document, roster: List[Player], wildcards: List[str], wildcard_history: Optional[Dict[str, str]] = None) -> None:
    document.add_heading("Full Roster", level=2)
    table = document.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["#", "Player", "DR", "Rating", "Season Record", "Singles", "Doubles", "2026 Status"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE)
        set_cell_shading(table.rows[0].cells[idx], NAVY)

    for index, player in enumerate(sorted(roster, key=lambda p: (p.dr is not None, p.dr or 0), reverse=True), start=1):
        row = table.add_row().cells
        if player.name in wildcards:
            prior = (wildcard_history or {}).get(player.name, "")
            status = f"NOT YET PLAYED ★" + (f" | {prior}" if prior else "")
        else:
            status = "Active"
        display_name = player.name
        if player.rating_type == "S":
            display_name += " ⚠"
        if player.dr is not None and player.dr >= 3.0:
            display_name = "◆ " + display_name

        values = [
            str(index),
            display_name,
            "-----" if player.dr is None else f"{player.dr:.2f}",
            player.rating_type,
            player.season_record,
            player.local_singles,
            player.local_doubles,
            status,
        ]
        for i, value in enumerate(values):
            set_cell_text(row[i], value, bold=(i == 1 and player.dr is not None and player.dr >= 3.0))

        if player.dr is not None and player.dr >= 3.0:
            for cell in row:
                set_cell_shading(cell, LIGHT_BLUE)
        if player.name in wildcards:
            for cell in row:
                set_cell_shading(cell, WARNING_YELLOW)
        if player.rating_type == "S":
            set_cell_shading(row[3], LOSS_RED)


def add_match_tables(document: Document, matches: List[MatchReport], team_label: str = "ETC") -> None:
    document.add_heading("Completed Matches", level=2)
    if not matches:
        document.add_paragraph(
            "No completed matches yet — every roster player is currently treated as a wildcard."
        )
        return

    for match in matches:
        document.add_paragraph(
            f"{match.date} vs {match.opponent} at {match.site} | Final: {match.team_name} {match.final_score}"
        )
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["Court", f"{team_label[:12]} Player", "DR", "Result", "Opponent Player", "DR"]
        for idx, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE)
            set_cell_shading(table.rows[0].cells[idx], NAVY)

        for court in match.courts:
            row = table.add_row().cells
            set_cell_text(row[0], court.court, bold=True)
            set_cell_text(row[1], format_player_line(court.team_players))
            set_cell_text(
                row[2],
                " / ".join("-----" if dr is None else f"{dr:.2f}" for _, dr, _, _ in court.team_players),
            )
            set_cell_text(row[3], f"{court.result} {court.score}", bold=True)
            set_cell_text(row[4], format_player_line(court.opponent_players))
            set_cell_text(
                row[5],
                " / ".join("-----" if dr is None else f"{dr:.2f}" for _, dr, _, _ in court.opponent_players),
            )
            set_cell_shading(row[3], WIN_GREEN if court.result == "W" else LOSS_RED)
        document.add_paragraph("")


def add_strategy_table(
    document: Document,
    strategy: Dict[str, CourtStrategy],
    wildcards: List[str],
    completed_matches: int,
    matches: List[MatchReport],
    roster: List[Player],
    team_name: str,
) -> None:
    document.add_heading("Strategy Notes", level=2)

    # --- Lineup Prediction header ---
    p = document.add_paragraph(f"Lineup Prediction — {team_name}")
    for run in p.runs:
        run.bold = True

    wins = sum(1 for m in matches if m.team_won_match)
    losses = len(matches) - wins
    p2 = document.add_paragraph(
        f"Based on {completed_matches} completed match{'es' if completed_matches != 1 else ''} ({wins}-{losses}). "
    )
    if completed_matches <= 4:
        p2.add_run("Small sample — treat as directional, not definitive.")

    document.add_paragraph()  # spacer

    # --- Overall Patterns bullets ---
    p_header = document.add_paragraph("Overall Patterns")
    for run in p_header.runs:
        run.bold = True

    court_wins = sum(
        1 for m in matches for c in m.courts if c.result == "W"
    )
    court_losses = sum(
        1 for m in matches for c in m.courts if c.result == "L"
    )
    total_courts = court_wins + court_losses

    bullet = document.add_paragraph(style="List Bullet")
    bullet.text = (
        f"Record: {wins}-{losses} ({court_wins}W-{court_losses}L by courts)."
    )

    # Singles summary
    s_wins = sum(
        1 for m in matches for c in m.courts
        if c.court.startswith("S") and c.result == "W"
    )
    s_losses = sum(
        1 for m in matches for c in m.courts
        if c.court.startswith("S") and c.result == "L"
    )
    if s_wins + s_losses > 0:
        if s_wins == 0:
            document.add_paragraph(
                f"Singles vulnerability: 0-{s_losses} across singles courts. Singles is their weakest area.",
                style="List Bullet",
            )
        elif s_losses == 0:
            document.add_paragraph(
                f"Singles strength: {s_wins}-0 across singles courts. Strong singles lineup.",
                style="List Bullet",
            )
        else:
            document.add_paragraph(
                f"Singles record: {s_wins}-{s_losses}.",
                style="List Bullet",
            )

    # Doubles summary
    d_wins = sum(
        1 for m in matches for c in m.courts
        if c.court.startswith("D") and c.result == "W"
    )
    d_losses = sum(
        1 for m in matches for c in m.courts
        if c.court.startswith("D") and c.result == "L"
    )
    if d_wins + d_losses > 0:
        if d_wins == 0:
            document.add_paragraph(
                f"Doubles vulnerability: 0-{d_losses} across doubles courts.",
                style="List Bullet",
            )
        elif d_losses == 0:
            document.add_paragraph(
                f"Doubles strength: {d_wins}-0 across doubles courts.",
                style="List Bullet",
            )
        else:
            document.add_paragraph(
                f"Doubles record: {d_wins}-{d_losses}.",
                style="List Bullet",
            )

    # Self-rated players on roster
    s_rated_roster = [p for p in roster if p.rating_type == "S"]
    if s_rated_roster:
        names = ", ".join(f"{p.name} ⚠ (S)" for p in s_rated_roster)
        document.add_paragraph(
            f"Self-rated flag: {names} — true level unknown, may play above listed rating.",
            style="List Bullet",
        )

    # Wildcards
    if wildcards:
        document.add_paragraph(
            f"Wildcards (never appeared): {', '.join(wildcards)} ★ — prepare for unknown lineup slots.",
            style="List Bullet",
        )
    else:
        document.add_paragraph(
            "No wildcards: All roster players have appeared in at least one match.",
            style="List Bullet",
        )

    # DR 3.0+ depth
    strong = [p for p in roster if p.dr is not None and p.dr >= 3.0]
    if strong:
        document.add_paragraph(
            f"DR ≥ 3.0 depth: {len(strong)} player{'s' if len(strong) != 1 else ''} "
            f"({', '.join(p.name for p in strong[:5])}).",
            style="List Bullet",
        )

    # --- 3-column strategy table ---
    table = document.add_table(rows=1, cols=3)
    table.style = "Table Grid"
    for idx, header in enumerate(["Court", f"{team_name.split('-')[0]} Likely Lineup", "Analysis"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE)
        set_cell_shading(table.rows[0].cells[idx], NAVY)

    for court in ["S1", "S2", "D1", "D2", "D3"]:
        cs = strategy[court]
        row = table.add_row().cells
        set_cell_text(row[0], court, bold=True)
        set_cell_text(row[1], cs.likely_lineup)
        set_cell_text(row[2], cs.analysis)

    # Disclaimer
    document.add_paragraph(
        f"Sample size is limited to {completed_matches} completed match{'es' if completed_matches != 1 else ''} "
        f"through {date.today().isoformat()}. Treat early lineup reads as directional rather than definitive."
    )


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def validate_report(document: Document, completed_matches: List[MatchReport], roster: List[Player]) -> None:
    if not roster:
        raise ValueError("Roster parsing failed; no players were found.")
    if not document.tables:
        raise ValueError("Generated report is missing tables.")


def build_output_name(year: int, level: str, team_name: str) -> str:
    return f"{year}_{slugify_filename(level)}_{slugify_filename(team_name)}_{report_date()}.docx"


def build_output_path(year: int, level: str, team_name: str, output_arg: Optional[str]) -> Path:
    reports_dir = Path(__file__).resolve().parent.parent.parent.parent / "reports"
    reports_dir.mkdir(parents=True, exist_ok=True)

    if output_arg:
        filename = Path(output_arg).name
        if not filename.lower().endswith(".docx"):
            filename += ".docx"
        return reports_dir / filename

    return reports_dir / build_output_name(year, level, team_name)


def parse_record(record: str) -> Tuple[int, int]:
    """Parse 'W-L' string into (wins, losses). Returns (0, 0) on failure."""
    m = re.match(r"(\d+)-(\d+)", record.strip())
    return (int(m.group(1)), int(m.group(2))) if m else (0, 0)


def fmt_record(wins: int, losses: int) -> str:
    return f"{wins}-{losses}"


def apply_manual_match_stats(roster: List[Player], manual_matches: List[MatchReport]) -> None:
    """
    Patch each roster player's local_singles, local_doubles, local_record,
    and season_record to include results from manually entered matches.
    Called after manual matches are loaded, before the document is built.
    """
    player_map = {canonicalize_name(p.name): p for p in roster}

    for match in manual_matches:
        for court in match.courts:
            is_singles = court.court.startswith("S")
            w_add = 1 if court.result == "W" else 0
            l_add = 1 if court.result == "L" else 0

            for name, _, _, _ in court.team_players:
                player = resolve_roster_player(name, player_map)
                if not player:
                    continue

                # Overall season record
                sw, sl = parse_record(player.season_record)
                player.season_record = fmt_record(sw + w_add, sl + l_add)

                # Local court-type record
                if is_singles:
                    lsw, lsl = parse_record(player.local_singles)
                    player.local_singles = fmt_record(lsw + w_add, lsl + l_add)
                else:
                    ldw, ldl = parse_record(player.local_doubles)
                    player.local_doubles = fmt_record(ldw + w_add, ldl + l_add)

                # Local overall record
                lrw, lrl = parse_record(player.local_record)
                player.local_record = fmt_record(lrw + w_add, lrl + l_add)


def load_manual_matches(
    path: str,
    team_name: str,
    rating_types: Dict[str, str],
    usta_ratings: Dict[str, str],
    roster: List[Player],
    year: int = 2026,
) -> Tuple[List[MatchReport], set]:
    """
    Load manually entered match data from a JSON file.
    Returns (list of MatchReport, set of team player names who appeared).
    JSON schema:
      [ { "date": "4/25/2026", "site": "...", "opponent": "...",
          "courts": [
            { "court": "S1", "team_players": ["Name"], "opponent_players": ["Name"],
              "score": "6-2 6-4", "result": "W" }, ... ] }, ... ]
    """
    roster_by_key: Dict[str, Player] = {canonicalize_name(p.name): p for p in roster}
    with open(path) as f:
        data = json.load(f)

    # Collect all unique opponent names and unresolved team names so we can fetch DRs in parallel.
    all_opp_names: set = set()
    unresolved_team_names: set = set()
    for m in data:
        for c in m["courts"]:
            for name in c["opponent_players"]:
                all_opp_names.add(name)
            for name in c["team_players"]:
                if resolve_roster_player(name, roster_by_key) is None:
                    unresolved_team_names.add(name)

    opp_info: Dict[str, Tuple[Optional[float], str, str]] = {}
    team_fallback_info: Dict[str, Tuple[Optional[float], str, str]] = {}
    with ThreadPoolExecutor(max_workers=10) as pool:
        future_to_opp = {pool.submit(fetch_player_info_by_name, name, year): name for name in all_opp_names}
        future_to_team = {pool.submit(fetch_player_info_by_name, name, year): name for name in unresolved_team_names}
        for future in as_completed(future_to_opp):
            name = future_to_opp[future]
            try:
                opp_info[name] = future.result()
            except Exception:
                opp_info[name] = (None, "—", "—")
        for future in as_completed(future_to_team):
            name = future_to_team[future]
            try:
                team_fallback_info[name] = future.result()
            except Exception:
                team_fallback_info[name] = (None, "—", "—")

    matches: List[MatchReport] = []
    seen_names: set = set()

    for m in data:
        courts: List[MatchCourt] = []
        team_wins = 0
        team_losses = 0
        for c in m["courts"]:
            result = c["result"]
            if result == "W":
                team_wins += 1
            else:
                team_losses += 1

            team_players: List[Tuple[str, Optional[float], str, str]] = []
            for name in c["team_players"]:
                roster_player = resolve_roster_player(name, roster_by_key)
                if roster_player:
                    resolved_name = roster_player.name
                    seen_names.add(resolved_name)
                    dr = roster_player.dr
                    rt = rating_types.get(resolved_name, "—")
                    ur = usta_ratings.get(resolved_name, "—")
                    team_players.append((resolved_name, dr, rt, ur))
                    continue

                seen_names.add(name)
                dr, rt, ur = team_fallback_info.get(name, (None, "—", "—"))
                team_players.append((name, dr, rt, ur))

            opp_players: List[Tuple[str, Optional[float], str, str]] = [
                (name, *opp_info.get(name, (None, "—", "—"))) for name in c["opponent_players"]
            ]

            courts.append(MatchCourt(
                court=c["court"],
                team_players=team_players,
                opponent_players=opp_players,
                score=c["score"],
                result=result,
            ))

        matches.append(MatchReport(
            date=m["date"],
            site=m["site"],
            team_name=team_name,
            opponent=m["opponent"],
            final_score=f"{team_wins}-{team_losses}",
            team_won_match=team_wins > team_losses,
            courts=courts,
        ))

    return matches, seen_names


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a USTA scouting report from TennisRecord.")
    parser.add_argument("--team", required=True, help="Exact TennisRecord team name.")
    parser.add_argument("--year", type=int, default=date.today().year, help="League year. Defaults to current year.")
    parser.add_argument(
        "--s",
        type=int,
        default=None,
        dest="suffix",
        help="Force &s=N disambiguator. Use when a team name matches multiple leagues (e.g. 40+ and 18+). Skips auto-detection.",
    )
    parser.add_argument(
        "--output",
        help="Optional output filename. Reports are always written into the repo-root reports/ folder.",
    )
    parser.add_argument(
        "--manual-matches",
        dest="manual_matches",
        help="Path to a JSON file of manually entered match results to include alongside scraped data.",
    )
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    league_label, flight_label, resolved_team_name, roster, match_rows = parse_team_page(args.team, args.year, forced_suffix=args.suffix)
    team_level = infer_team_level(league_label, roster)

    profile_links: Dict[str, str] = {player.name: player.href for player in roster}
    completed_matches = []
    participant_names = set()
    participant_links: Dict[str, str] = {}

    for _, _, result, href in match_rows:
        if result == "0-0":
            continue
        full_url = urljoin(BASE_URL, href)
        html = fetch(full_url)
        msoup = BeautifulSoup(html, "lxml")
        for a in msoup.find_all("a", href=lambda h: h and "/adult/profile.aspx?playername=" in h):
            name = normalize_space(a.get_text(" ", strip=True))
            participant_names.add(name)
            participant_links[name] = a.get("href")

    profile_links.update(participant_links)
    player_ratings: Dict[str, Tuple[str, str]] = {}
    with ThreadPoolExecutor(max_workers=10) as pool:
        future_to_name = {pool.submit(parse_player_rating, href): name for name, href in profile_links.items()}
        for future in as_completed(future_to_name):
            name = future_to_name[future]
            try:
                player_ratings[name] = future.result()
            except Exception:
                player_ratings[name] = ("—", "—")
    rating_types = {name: rating_info[0] for name, rating_info in player_ratings.items()}
    usta_ratings = {name: rating_info[1] for name, rating_info in player_ratings.items()}

    for player in roster:
        player.rating_type = rating_types.get(player.name, "—")
        if usta_ratings.get(player.name, "—") == "—":
            usta_ratings[player.name] = f"{player.ntrp}{player.rating_type}"

    seen_hrefs = set()
    for _, _, result, href in match_rows:
        if result == "0-0":
            continue
        if href in seen_hrefs:
            continue
        seen_hrefs.add(href)
        completed_matches.append(parse_match(href, rating_types, usta_ratings, resolved_team_name))

    if args.manual_matches:
        manual_matches, manual_names = load_manual_matches(
            args.manual_matches, resolved_team_name, rating_types, usta_ratings, roster, year=args.year
        )
        completed_matches.extend(manual_matches)
        participant_names.update(manual_names)
        apply_manual_match_stats(roster, manual_matches)

    wildcards = sorted(player.name for player in roster if player.name not in participant_names)

    # Enrich wildcards with prior-season history
    wildcard_players = {p.name: p for p in roster if p.name in wildcards}
    wildcard_history: Dict[str, str] = {}
    wc_with_href = [(name, player) for name, player in wildcard_players.items() if player.href]
    with ThreadPoolExecutor(max_workers=8) as pool:
        future_to_wc = {pool.submit(fetch_wildcard_prior_season, player.href, args.year): name
                        for name, player in wc_with_href}
        for future in as_completed(future_to_wc):
            name = future_to_wc[future]
            try:
                history = future.result()
                if history:
                    wildcard_history[name] = history
            except Exception:
                pass

    strategy = build_strategy(completed_matches, wildcards, roster, resolved_team_name)
    wins = sum(1 for match in completed_matches if match.team_won_match)
    losses = len(completed_matches) - wins
    most_recent_match = completed_matches[-1].date if completed_matches else None

    # Derive a short team label for column headers
    team_label = resolved_team_name

    document = Document()
    configure_document(document)
    add_title(document, resolved_team_name, league_label, flight_label, f"{wins}-{losses}", most_recent_match=most_recent_match)
    add_legend(document, team_name=resolved_team_name)
    add_roster_table(document, roster, wildcards, wildcard_history=wildcard_history)
    add_match_tables(document, completed_matches, team_label=team_label)
    add_strategy_table(document, strategy, wildcards, len(completed_matches), completed_matches, roster, resolved_team_name)
    validate_report(document, completed_matches, roster)

    output_path = build_output_path(args.year, team_level, resolved_team_name, args.output)
    document.save(output_path)
    print(output_path)
    print(f"completed_matches={len(completed_matches)}")
    print(f"wildcards={len(wildcards)}")


if __name__ == "__main__":
    main()
