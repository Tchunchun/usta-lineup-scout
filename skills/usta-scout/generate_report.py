import argparse
import re
import subprocess
from dataclasses import dataclass
from datetime import date
from pathlib import Path
from typing import Dict, List, Optional, Tuple
from urllib.parse import quote, urljoin

from bs4 import BeautifulSoup
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor


BASE_URL = "https://www.tennisrecord.com"

NAVY = "1F3864"
LIGHT_BLUE = "D6E4F0"
MID_BLUE = "2E75B6"
WIN_GREEN = "E2EFDA"
LOSS_RED = "FCE4D6"
WARNING_YELLOW = "FFF2CC"
WHITE = "FFFFFF"


@dataclass
class Player:
    name: str
    href: str
    location: str
    ntrp: str
    season_record: str
    local_singles: str
    local_doubles: str
    local_record: str
    dr: Optional[float]
    rating_type: str = "—"


@dataclass
class MatchCourt:
    court: str
    team_players: List[Tuple[str, Optional[float], str]]
    opponent_players: List[Tuple[str, Optional[float], str]]
    score: str
    result: str


@dataclass
class MatchReport:
    date: str
    site: str
    team_name: str
    opponent: str
    final_score: str
    team_won_match: bool
    courts: List[MatchCourt]


def fetch(url: str) -> str:
    safe_url = quote(url, safe=":/?&=%")
    return subprocess.check_output(["curl", "-L", "--silent", safe_url], text=True)


def soup_from_url(url: str) -> BeautifulSoup:
    return BeautifulSoup(fetch(url), "lxml")


def normalize_space(text: str) -> str:
    return " ".join(text.split())


def slugify_filename(value: str) -> str:
    slug = re.sub(r"[^A-Za-z0-9]+", "_", value).strip("_")
    return slug or "Scouting_Report"


def build_team_url(team_name: str, year: int, suffix: Optional[int] = None) -> str:
    url = f"{BASE_URL}/adult/teamprofile.aspx?year={year}&teamname={quote(team_name, safe='')}"
    if suffix is not None:
        url += f"&s={suffix}"
    return url


def set_cell_shading(cell, fill: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:fill"), fill)
    tc_pr.append(shd)


def set_cell_text(cell, text: str, bold: bool = False, color: Optional[str] = None) -> None:
    cell.text = ""
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = p.add_run(text)
    run.bold = bold
    if color:
        run.font.color.rgb = RGBColor.from_string(color)
    run.font.size = Pt(9)


def parse_rating_type(profile_href: str) -> str:
    url = urljoin(BASE_URL, profile_href)
    html = fetch(url)
    match = re.search(r"<span style=\"font-weight:bold;\">(\d\.\d)\s+([A-Z])</span>", html)
    if match:
        rating_type = match.group(2)
        return rating_type if rating_type in {"C", "S"} else "—"
    return "—"


def parse_dr(value: str) -> Optional[float]:
    try:
        return float(value)
    except ValueError:
        return None


def parse_players_from_cell(cell) -> List[Tuple[str, Optional[float]]]:
    players: List[Tuple[str, Optional[float]]] = []
    for anchor in cell.find_all("a", href=lambda h: h and "/adult/profile.aspx?playername=" in h):
        name = normalize_space(anchor.get_text(" ", strip=True))
        trailing = anchor.next_sibling or ""
        rating_match = re.search(r"\(([-0-9.]+|-----)\)", str(trailing))
        dr = None if not rating_match or rating_match.group(1) == "-----" else float(rating_match.group(1))
        players.append((name, dr))
    return players


def format_player_line(players: List[Tuple[str, Optional[float], str]]) -> str:
    formatted = []
    for name, dr, rating_type in players:
        dr_text = "-----" if dr is None else f"{dr:.2f}"
        formatted.append(f"{name} ({rating_type}) [{dr_text}]")
    return " / ".join(formatted)


def parse_team_page(team_name: str, year: int) -> Tuple[str, str, str, List[Player], List[Tuple[str, str, str, str]]]:
    soup = None
    for suffix in [None, 1, 2, 3, 4, 5]:
        candidate = soup_from_url(build_team_url(team_name, year, suffix))
        tables = candidate.find_all("table")
        if len(tables) < 5:
            continue
        meta_rows = tables[1].find_all("tr")
        if len(meta_rows) < 3:
            continue
        resolved_name = normalize_space(meta_rows[2].get_text(" ", strip=True))
        if resolved_name == team_name:
            soup = candidate
            break
    if soup is None:
        raise ValueError(f"Unable to locate team page for '{team_name}' in {year}.")

    tables = soup.find_all("table")
    meta_rows = tables[1].find_all("tr")
    league_label = normalize_space(meta_rows[0].get_text(" ", strip=True))
    flight_label = normalize_space(meta_rows[1].get_text(" ", strip=True))
    resolved_team_name = normalize_space(meta_rows[2].get_text(" ", strip=True))

    roster: List[Player] = []
    roster_table = tables[2]
    for tr in roster_table.find_all("tr")[1:]:
        cells = tr.find_all("td")
        if len(cells) != 9:
            continue
        link = tr.find("a", href=lambda h: h and "/adult/profile.aspx?playername=" in h)
        if not link:
            continue
        texts = [normalize_space(td.get_text(" ", strip=True)) for td in cells]
        roster.append(
            Player(
                name=normalize_space(link.get_text(" ", strip=True)),
                href=link.get("href"),
                location=texts[1],
                ntrp=texts[2],
                season_record=texts[3],
                local_singles=texts[4],
                local_doubles=texts[5],
                local_record=texts[6],
                dr=parse_dr(texts[7]),
            )
        )

    schedule_table = tables[4]
    matches: List[Tuple[str, str, str, str]] = []
    for tr in schedule_table.find_all("tr")[1:]:
        cells = tr.find_all("td")
        if len(cells) != 5:
            continue
        match_date = normalize_space(cells[0].get_text(" ", strip=True))
        opponent = normalize_space(cells[2].get_text(" ", strip=True))
        result = normalize_space(cells[4].get_text(" ", strip=True))
        link = cells[4].find("a", href=lambda h: h and "/adult/matchresults.aspx?year=" in h)
        if link:
            matches.append((match_date, opponent, result, link.get("href")))
    return league_label, flight_label, resolved_team_name, roster, matches


def parse_match(match_href: str, rating_types: Dict[str, str], team_name: str) -> MatchReport:
    soup = soup_from_url(urljoin(BASE_URL, match_href))
    tables = soup.find_all("table")

    details_text = normalize_space(tables[1].get_text(" ", strip=True))
    date_match = re.search(r"Scheduled Date:\s*([0-9/]+)", details_text)
    site_match = re.search(r"Match Site:\s*(.*?)\s*Match Criteria:", details_text)

    summary_rows = tables[2].find_all("tr")[1:]
    summary = []
    for row in summary_rows:
        cells = row.find_all("td")
        if len(cells) >= 4:
            summary.append([normalize_space(cell.get_text(" ", strip=True)) for cell in cells])
    team_is_left = summary[0][0] == team_name
    opponent = summary[1][0] if team_is_left else summary[0][0]
    wins = summary[0][1] if team_is_left else summary[1][1]
    losses = summary[1][1] if team_is_left else summary[0][1]
    team_won_match = int(wins) > int(losses)

    court_labels = ["Singles #1", "Singles #2", "Doubles #1", "Doubles #2", "Doubles #3"]
    court_keys = ["S1", "S2", "D1", "D2", "D3"]
    courts: List[MatchCourt] = []

    for label, key in zip(court_labels, court_keys):
        marker = soup.find(string=lambda s: isinstance(s, str) and label in s)
        if not marker:
            continue
        court_table = marker.parent.find_next("table")
        rows = court_table.find_all("tr")
        if len(rows) < 2:
            continue
        row = rows[1]
        cells = row.find_all("td")
        left_players = parse_players_from_cell(cells[0])
        right_players = parse_players_from_cell(cells[-1])
        score = normalize_space(cells[3].get_text(" ", strip=True)).replace(" - ", "-")
        row_html = str(row)
        winner_side = "left" if "arrowhead_right.png" in row_html else "right"
        team_side = "left" if team_is_left else "right"
        result = "W" if winner_side == team_side else "L"

        team_players_raw = left_players if team_side == "left" else right_players
        opp_players_raw = right_players if team_side == "left" else left_players
        team_players = [(name, dr, rating_types.get(name, "—")) for name, dr in team_players_raw]
        opponent_players = [(name, dr, rating_types.get(name, "—")) for name, dr in opp_players_raw]

        courts.append(
            MatchCourt(
                court=key,
                team_players=team_players,
                opponent_players=opponent_players,
                score=score,
                result=result,
            )
        )

    return MatchReport(
        date=date_match.group(1) if date_match else "",
        site=site_match.group(1) if site_match else "",
        team_name=team_name,
        opponent=opponent,
        final_score=f"{wins}-{losses}",
        team_won_match=team_won_match,
        courts=courts,
    )


def build_strategy(matches: List[MatchReport], wildcards: List[str]) -> Dict[str, str]:
    strategy: Dict[str, str] = {}
    court_map: Dict[str, List[MatchCourt]] = {"S1": [], "S2": [], "D1": [], "D2": [], "D3": []}
    for match in matches:
        for court in match.courts:
            court_map[court.court].append(court)

    for court, entries in court_map.items():
        if not entries:
            strategy[court] = "No completed sample yet."
            continue
        opponents = []
        unknowns = []
        for entry in entries:
            names = ", ".join(
                f"{name} ({'-----' if dr is None else f'{dr:.2f}'}, {rating_type})"
                for name, dr, rating_type in entry.opponent_players
            )
            opponents.append(f"{names} in a {entry.result} ({entry.score})")
            for name, dr, rating_type in entry.opponent_players:
                if rating_type == "S" or dr is None:
                    unknowns.append(name)
        note = " ; ".join(opponents)
        recommendation = "Keep pressure on this line and make them prove depth."
        if court.startswith("S"):
            recommendation = "Favor steady first-strike singles patterns and make them hit through you."
        if court.startswith("D"):
            recommendation = "Test team chemistry early with poaches, lobs, and serve-return pressure."
        if unknowns:
            recommendation += f" Treat {', '.join(sorted(set(unknowns)))} as higher-variance opponents."
        if wildcards:
            recommendation += " Keep one flexible pairing ready for roster surprises."
        strategy[court] = f"{note}. {recommendation}"
    return strategy


def add_title(document: Document, team_name: str, league_label: str, flight_label: str, season_record: str) -> None:
    p = document.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = p.add_run(f"{team_name} Scouting Report")
    run.bold = True
    run.font.size = Pt(18)
    run.font.color.rgb = RGBColor.from_string(NAVY)

    for line in [
        f"League: {league_label}",
        f"Flight: {flight_label}",
        f"Prepared: {date.today().isoformat()}",
        f"Completed Match Record: {season_record}",
    ]:
        p = document.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run(line)
        r.font.size = Pt(10)


def add_legend(document: Document) -> None:
    document.add_heading("Legend", level=2)
    for line in [
        "◆ = DR 3.0+ player",
        "⚠ = Self-rated player (S)",
        "★ = Player on roster with no completed-match appearance yet",
        "Green result cells = ETC win",
        "Red result cells = ETC loss",
    ]:
        p = document.add_paragraph(style="List Bullet")
        p.add_run(line)


def add_roster_table(document: Document, roster: List[Player], wildcards: List[str]) -> None:
    document.add_heading("Full Roster", level=2)
    table = document.add_table(rows=1, cols=8)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.style = "Table Grid"
    headers = ["#", "Player", "DR", "Rating", "Season Record", "Singles", "Doubles", "2026 Status"]
    for idx, header in enumerate(headers):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE)
        set_cell_shading(table.rows[0].cells[idx], NAVY)

    for index, player in enumerate(sorted(roster, key=lambda p: (p.dr is not None, p.dr or 0), reverse=True), start=1):
        row = table.add_row().cells
        status = "NOT YET PLAYED ★" if player.name in wildcards else "Active"
        display_name = player.name
        if player.rating_type == "S":
            display_name += " ⚠"
        if player.dr is not None and player.dr >= 3.0:
            display_name = "◆ " + display_name

        values = [
            str(index),
            display_name,
            "-----" if player.dr is None else f"{player.dr:.2f}",
            player.rating_type,
            player.season_record,
            player.local_singles,
            player.local_doubles,
            status,
        ]
        for i, value in enumerate(values):
            set_cell_text(row[i], value, bold=(i == 1 and player.dr is not None and player.dr >= 3.0))

        if player.dr is not None and player.dr >= 3.0:
            for cell in row:
                set_cell_shading(cell, LIGHT_BLUE)
        if player.name in wildcards:
            for cell in row:
                set_cell_shading(cell, WARNING_YELLOW)
        if player.rating_type == "S":
            set_cell_shading(row[3], LOSS_RED)


def add_match_tables(document: Document, matches: List[MatchReport]) -> None:
    document.add_heading("Completed Matches", level=2)
    for match in matches:
        document.add_paragraph(
            f"{match.date} vs {match.opponent} at {match.site} | Final: {match.team_name} {match.final_score}"
        )
        table = document.add_table(rows=1, cols=6)
        table.style = "Table Grid"
        headers = ["Court", "ETC Player", "DR", "Result", "Opponent Player", "DR"]
        for idx, header in enumerate(headers):
            set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE)
            set_cell_shading(table.rows[0].cells[idx], NAVY)

        for court in match.courts:
            row = table.add_row().cells
            set_cell_text(row[0], court.court, bold=True)
            set_cell_text(row[1], format_player_line(court.team_players))
            set_cell_text(
                row[2],
                " / ".join("-----" if dr is None else f"{dr:.2f}" for _, dr, _ in court.team_players),
            )
            set_cell_text(row[3], f"{court.result} {court.score}", bold=True)
            set_cell_text(row[4], format_player_line(court.opponent_players))
            set_cell_text(
                row[5],
                " / ".join("-----" if dr is None else f"{dr:.2f}" for _, dr, _ in court.opponent_players),
            )
            set_cell_shading(row[3], WIN_GREEN if court.result == "W" else LOSS_RED)
        document.add_paragraph("")


def add_strategy_table(document: Document, strategy: Dict[str, str], wildcards: List[str], completed_matches: int) -> None:
    document.add_heading("Strategy Notes", level=2)
    intro = (
        f"Sample size is limited to {completed_matches} completed matches through {date.today().isoformat()}. "
        "Treat early lineup reads as directional rather than definitive."
    )
    document.add_paragraph(intro)
    if wildcards:
        document.add_paragraph("Wildcards: " + ", ".join(wildcards))

    table = document.add_table(rows=1, cols=2)
    table.style = "Table Grid"
    for idx, header in enumerate(["Court", "Recommendation"]):
        set_cell_text(table.rows[0].cells[idx], header, bold=True, color=WHITE)
        set_cell_shading(table.rows[0].cells[idx], NAVY)

    for court in ["S1", "S2", "D1", "D2", "D3"]:
        row = table.add_row().cells
        set_cell_text(row[0], court, bold=True)
        set_cell_text(row[1], strategy[court])


def configure_document(document: Document) -> None:
    section = document.sections[0]
    section.top_margin = Inches(0.6)
    section.bottom_margin = Inches(0.6)
    section.left_margin = Inches(0.6)
    section.right_margin = Inches(0.6)


def validate_report(document: Document, completed_matches: List[MatchReport], roster: List[Player]) -> None:
    if not roster:
        raise ValueError("Roster parsing failed; no players were found.")
    if not completed_matches:
        raise ValueError("No completed matches were parsed.")
    if not document.tables:
        raise ValueError("Generated report is missing tables.")


def build_output_name(team_name: str) -> str:
    return f"{slugify_filename(team_name)}_Scouting_Report.docx"


def parse_args() -> argparse.Namespace:
    parser = argparse.ArgumentParser(description="Generate a USTA scouting report from TennisRecord.")
    parser.add_argument("--team", required=True, help="Exact TennisRecord team name.")
    parser.add_argument("--year", type=int, default=date.today().year, help="League year. Defaults to current year.")
    parser.add_argument("--output", help="Optional output .docx path.")
    return parser.parse_args()


def main() -> None:
    args = parse_args()
    league_label, flight_label, resolved_team_name, roster, match_rows = parse_team_page(args.team, args.year)

    profile_links: Dict[str, str] = {player.name: player.href for player in roster}
    completed_matches = []
    participant_names = set()
    participant_links: Dict[str, str] = {}

    for _, _, result, href in match_rows:
        if result == "0-0":
            continue
        full_url = urljoin(BASE_URL, href)
        html = fetch(full_url)
        msoup = BeautifulSoup(html, "lxml")
        for a in msoup.find_all("a", href=lambda h: h and "/adult/profile.aspx?playername=" in h):
            name = normalize_space(a.get_text(" ", strip=True))
            participant_names.add(name)
            participant_links[name] = a.get("href")

    profile_links.update(participant_links)
    rating_types = {name: parse_rating_type(href) for name, href in profile_links.items()}

    for player in roster:
        player.rating_type = rating_types.get(player.name, "—")

    seen_hrefs = set()
    for _, _, result, href in match_rows:
        if result == "0-0":
            continue
        if href in seen_hrefs:
            continue
        seen_hrefs.add(href)
        completed_matches.append(parse_match(href, rating_types, resolved_team_name))

    wildcards = sorted(player.name for player in roster if player.name not in participant_names)
    strategy = build_strategy(completed_matches, wildcards)
    wins = sum(1 for match in completed_matches if match.team_won_match)
    losses = len(completed_matches) - wins

    document = Document()
    configure_document(document)
    add_title(document, resolved_team_name, league_label, flight_label, f"{wins}-{losses}")
    add_legend(document)
    add_roster_table(document, roster, wildcards)
    add_match_tables(document, completed_matches)
    add_strategy_table(document, strategy, wildcards, len(completed_matches))
    validate_report(document, completed_matches, roster)

    output_path = Path(args.output) if args.output else Path.cwd() / build_output_name(resolved_team_name)
    document.save(output_path)
    print(output_path)
    print(f"completed_matches={len(completed_matches)}")
    print(f"wildcards={len(wildcards)}")


if __name__ == "__main__":
    main()
